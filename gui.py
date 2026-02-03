#!/usr/bin/env python3
"""
Interface graphique pour la transcription audio avec Whisper.
"""

import threading
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk
from pathlib import Path

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    DND_AVAILABLE = True
except ImportError:
    DND_AVAILABLE = False

from transcribe import transcribe_audio

# Import optionnel pour l'export DOCX
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class SettingsDialog(tk.Toplevel):
    """Fenêtre de paramètres."""

    def __init__(self, parent, current_model, current_language):
        super().__init__(parent)
        self.title("Paramètres")
        self.geometry("350x250")
        self.resizable(False, False)
        self.transient(parent)
        self.grab_set()

        self.result_model = current_model
        self.result_language = current_language

        # Frame principal
        main_frame = ttk.Frame(self, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Sélection du modèle
        ttk.Label(main_frame, text="Modèle Whisper:", font=("", 10, "bold")).pack(anchor=tk.W)

        models_info = [
            ("tiny", "Tiny - Rapide (~1GB VRAM)"),
            ("base", "Base - Équilibré (~1GB VRAM)"),
            ("small", "Small - Bon (~2GB VRAM)"),
            ("medium", "Medium - Haute qualité (~5GB VRAM)"),
            ("large", "Large - Meilleure qualité (~10GB VRAM)"),
        ]

        self.model_var = tk.StringVar(value=current_model)

        model_frame = ttk.Frame(main_frame)
        model_frame.pack(fill=tk.X, pady=(5, 15))

        for model_id, model_label in models_info:
            rb = ttk.Radiobutton(
                model_frame,
                text=model_label,
                value=model_id,
                variable=self.model_var
            )
            rb.pack(anchor=tk.W, pady=2)

        # Sélection de la langue
        ttk.Label(main_frame, text="Langue:", font=("", 10, "bold")).pack(anchor=tk.W, pady=(10, 0))

        lang_frame = ttk.Frame(main_frame)
        lang_frame.pack(fill=tk.X, pady=5)

        self.language_var = tk.StringVar(value=current_language or "auto")

        languages = [
            ("auto", "Détection automatique"),
            ("fr", "Français"),
            ("en", "Anglais"),
            ("es", "Espagnol"),
            ("de", "Allemand"),
            ("it", "Italien"),
        ]

        self.lang_combo = ttk.Combobox(
            lang_frame,
            textvariable=self.language_var,
            values=[f"{code} - {name}" for code, name in languages],
            state="readonly",
            width=30
        )
        self.lang_combo.pack(fill=tk.X)

        # Trouver l'index correspondant
        for i, (code, _) in enumerate(languages):
            if code == (current_language or "auto"):
                self.lang_combo.current(i)
                break
        else:
            self.lang_combo.current(0)

        # Boutons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=(20, 0))

        ttk.Button(btn_frame, text="Annuler", command=self.cancel).pack(side=tk.RIGHT, padx=(5, 0))
        ttk.Button(btn_frame, text="Appliquer", command=self.apply).pack(side=tk.RIGHT)

        # Centrer la fenêtre
        self.update_idletasks()
        x = parent.winfo_x() + (parent.winfo_width() - self.winfo_width()) // 2
        y = parent.winfo_y() + (parent.winfo_height() - self.winfo_height()) // 2
        self.geometry(f"+{x}+{y}")

    def apply(self):
        self.result_model = self.model_var.get()
        lang_selection = self.language_var.get()
        lang_code = lang_selection.split(" - ")[0] if " - " in lang_selection else lang_selection
        self.result_language = None if lang_code == "auto" else lang_code
        self.destroy()

    def cancel(self):
        self.destroy()


class TranscriptionApp:
    """Application principale de transcription."""

    VALID_EXTENSIONS = {".mp3", ".wav", ".m4a", ".flac", ".ogg", ".webm", ".mp4"}

    def __init__(self):
        # Créer la fenêtre principale avec support drag & drop si disponible
        if DND_AVAILABLE:
            self.root = TkinterDnD.Tk()
        else:
            self.root = tk.Tk()

        self.root.title("Transcription Audio - Whisper")
        self.root.geometry("700x600")
        self.root.minsize(500, 400)

        # Variables
        self.current_file = None
        self.model = "medium"  # Modèle par défaut: medium
        self.language = None   # Détection auto par défaut
        self.is_transcribing = False

        self.setup_ui()
        self.setup_styles()

    def setup_styles(self):
        """Configure les styles ttk."""
        style = ttk.Style()
        style.configure("Drop.TFrame", background="#e8f4f8")
        style.configure("DropHover.TFrame", background="#c8e6f0")
        style.configure("Title.TLabel", font=("", 12, "bold"))
        style.configure("Info.TLabel", foreground="#666666")
        style.configure("Success.TLabel", foreground="#2e7d32")
        style.configure("Action.TButton", font=("", 10))

    def setup_ui(self):
        """Configure l'interface utilisateur."""
        # Frame principal
        main_frame = ttk.Frame(self.root, padding="15")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Header avec titre et bouton paramètres
        header_frame = ttk.Frame(main_frame)
        header_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(
            header_frame,
            text="Transcription Audio",
            style="Title.TLabel"
        ).pack(side=tk.LEFT)

        self.settings_btn = ttk.Button(
            header_frame,
            text="Paramètres",
            command=self.open_settings
        )
        self.settings_btn.pack(side=tk.RIGHT)

        # Info modèle actuel
        self.model_info_label = ttk.Label(
            header_frame,
            text=f"Modèle: {self.model}",
            style="Info.TLabel"
        )
        self.model_info_label.pack(side=tk.RIGHT, padx=(0, 15))

        # Zone de drop
        self.drop_frame = ttk.Frame(main_frame, style="Drop.TFrame", padding="30")
        self.drop_frame.pack(fill=tk.X, pady=(0, 10))

        # Contenu de la zone de drop
        drop_content = ttk.Frame(self.drop_frame, style="Drop.TFrame")
        drop_content.pack(expand=True)

        self.drop_icon_label = ttk.Label(
            drop_content,
            text="📁",
            font=("", 36),
            style="Drop.TLabel"
        )
        self.drop_icon_label.pack()

        self.drop_label = ttk.Label(
            drop_content,
            text="Glissez-déposez un fichier audio ici\nou cliquez pour sélectionner",
            justify=tk.CENTER,
            style="Drop.TLabel"
        )
        self.drop_label.pack(pady=(10, 0))

        # Configurer le drag & drop si disponible
        if DND_AVAILABLE:
            self.drop_frame.drop_target_register(DND_FILES)
            self.drop_frame.dnd_bind("<<Drop>>", self.on_drop)
            self.drop_frame.dnd_bind("<<DragEnter>>", self.on_drag_enter)
            self.drop_frame.dnd_bind("<<DragLeave>>", self.on_drag_leave)
        else:
            self.drop_label.configure(
                text="Cliquez pour sélectionner un fichier audio\n(Installez tkinterdnd2 pour le glisser-déposer)"
            )

        # Rendre la zone cliquable
        for widget in [self.drop_frame, drop_content, self.drop_icon_label, self.drop_label]:
            widget.bind("<Button-1>", self.select_file)
            widget.configure(cursor="hand2")

        # Fichier sélectionné
        self.file_frame = ttk.Frame(main_frame)
        self.file_frame.pack(fill=tk.X, pady=(0, 10))

        self.file_label = ttk.Label(
            self.file_frame,
            text="Aucun fichier sélectionné",
            style="Info.TLabel"
        )
        self.file_label.pack(side=tk.LEFT)

        self.clear_btn = ttk.Button(
            self.file_frame,
            text="✕",
            width=3,
            command=self.clear_file
        )
        # Caché par défaut

        # Bouton de transcription
        self.transcribe_btn = ttk.Button(
            main_frame,
            text="Lancer la transcription",
            command=self.start_transcription,
            style="Action.TButton",
            state=tk.DISABLED
        )
        self.transcribe_btn.pack(fill=tk.X, pady=(0, 10))

        # Barre de progression
        self.progress = ttk.Progressbar(main_frame, mode="indeterminate")
        self.progress.pack(fill=tk.X, pady=(0, 10))

        # Status
        self.status_label = ttk.Label(main_frame, text="", style="Info.TLabel")
        self.status_label.pack(fill=tk.X, pady=(0, 10))

        # Zone de résultat
        result_frame = ttk.LabelFrame(main_frame, text="Transcription", padding="10")
        result_frame.pack(fill=tk.BOTH, expand=True)

        self.result_text = scrolledtext.ScrolledText(
            result_frame,
            wrap=tk.WORD,
            font=("", 10),
            state=tk.DISABLED
        )
        self.result_text.pack(fill=tk.BOTH, expand=True)

        # Boutons de résultat
        result_btns = ttk.Frame(result_frame)
        result_btns.pack(fill=tk.X, pady=(10, 0))

        self.copy_btn = ttk.Button(
            result_btns,
            text="Copier",
            command=self.copy_result,
            state=tk.DISABLED
        )
        self.copy_btn.pack(side=tk.LEFT)

        self.save_btn = ttk.Button(
            result_btns,
            text="Sauvegarder",
            command=self.save_result,
            state=tk.DISABLED
        )
        self.save_btn.pack(side=tk.LEFT, padx=(5, 0))

    def on_drop(self, event):
        """Gère le drop d'un fichier."""
        self.drop_frame.configure(style="Drop.TFrame")

        # Nettoyer le chemin (enlever les accolades sur certains OS)
        file_path = event.data.strip()
        if file_path.startswith("{") and file_path.endswith("}"):
            file_path = file_path[1:-1]

        self.set_file(file_path)

    def on_drag_enter(self, event):
        """Effet visuel quand un fichier entre dans la zone."""
        self.drop_frame.configure(style="DropHover.TFrame")

    def on_drag_leave(self, event):
        """Effet visuel quand un fichier quitte la zone."""
        self.drop_frame.configure(style="Drop.TFrame")

    def select_file(self, event=None):
        """Ouvre le dialogue de sélection de fichier."""
        if self.is_transcribing:
            return

        filetypes = [
            ("Fichiers audio", "*.mp3 *.wav *.m4a *.flac *.ogg *.webm *.mp4"),
            ("Tous les fichiers", "*.*")
        ]
        file_path = filedialog.askopenfilename(filetypes=filetypes)
        if file_path:
            self.set_file(file_path)

    def set_file(self, file_path):
        """Définit le fichier à transcrire."""
        path = Path(file_path)

        if not path.exists():
            messagebox.showerror("Erreur", f"Le fichier n'existe pas:\n{file_path}")
            return

        if path.suffix.lower() not in self.VALID_EXTENSIONS:
            messagebox.showwarning(
                "Attention",
                f"Extension '{path.suffix}' non standard.\n"
                f"Extensions supportées: {', '.join(self.VALID_EXTENSIONS)}"
            )

        self.current_file = path
        self.file_label.configure(text=f"📄 {path.name}", style="Success.TLabel")
        self.clear_btn.pack(side=tk.RIGHT)
        self.transcribe_btn.configure(state=tk.NORMAL)
        self.status_label.configure(text="")

    def clear_file(self):
        """Efface le fichier sélectionné."""
        self.current_file = None
        self.file_label.configure(text="Aucun fichier sélectionné", style="Info.TLabel")
        self.clear_btn.pack_forget()
        self.transcribe_btn.configure(state=tk.DISABLED)

    def open_settings(self):
        """Ouvre la fenêtre de paramètres."""
        dialog = SettingsDialog(self.root, self.model, self.language)
        self.root.wait_window(dialog)

        self.model = dialog.result_model
        self.language = dialog.result_language

        lang_text = self.language if self.language else "auto"
        self.model_info_label.configure(text=f"Modèle: {self.model} | Langue: {lang_text}")

    def start_transcription(self):
        """Lance la transcription dans un thread séparé."""
        if not self.current_file or self.is_transcribing:
            return

        self.is_transcribing = True
        self.transcribe_btn.configure(state=tk.DISABLED)
        self.settings_btn.configure(state=tk.DISABLED)
        self.progress.start(10)
        self.status_label.configure(text="Chargement du modèle...")

        # Effacer le résultat précédent
        self.result_text.configure(state=tk.NORMAL)
        self.result_text.delete(1.0, tk.END)
        self.result_text.configure(state=tk.DISABLED)

        # Lancer dans un thread
        thread = threading.Thread(target=self.do_transcription, daemon=True)
        thread.start()

    def do_transcription(self):
        """Effectue la transcription (dans un thread)."""
        try:
            self.update_status("Chargement du modèle (peut prendre du temps)...")

            result = transcribe_audio(
                str(self.current_file),
                model_name=self.model,
                language=self.language
            )

            transcript = result.get("text", "").strip()
            if not transcript:
                raise ValueError("La transcription est vide ou a échoué")
            detected_lang = result.get("language", "inconnu")

            # Mettre à jour l'interface depuis le thread principal
            self.root.after(0, lambda: self.show_result(transcript, detected_lang))

        except Exception as e:
            self.root.after(0, lambda: self.show_error(str(e)))

    def update_status(self, text):
        """Met à jour le status depuis un thread."""
        self.root.after(0, lambda: self.status_label.configure(text=text))

    def show_result(self, transcript, detected_language):
        """Affiche le résultat de la transcription."""
        self.is_transcribing = False
        self.progress.stop()
        self.transcribe_btn.configure(state=tk.NORMAL)
        self.settings_btn.configure(state=tk.NORMAL)

        self.status_label.configure(
            text=f"Transcription terminée ! Langue détectée: {detected_language}",
            style="Success.TLabel"
        )

        self.result_text.configure(state=tk.NORMAL)
        self.result_text.delete(1.0, tk.END)
        self.result_text.insert(1.0, transcript)
        self.result_text.configure(state=tk.DISABLED)

        self.copy_btn.configure(state=tk.NORMAL)
        self.save_btn.configure(state=tk.NORMAL)

    def show_error(self, error_msg):
        """Affiche une erreur."""
        self.is_transcribing = False
        self.progress.stop()
        self.transcribe_btn.configure(state=tk.NORMAL)
        self.settings_btn.configure(state=tk.NORMAL)
        self.status_label.configure(text=f"Erreur: {error_msg}")
        messagebox.showerror("Erreur de transcription", error_msg)

    def copy_result(self):
        """Copie le résultat dans le presse-papier."""
        self.root.clipboard_clear()
        self.root.clipboard_append(self.result_text.get(1.0, tk.END).strip())
        self.status_label.configure(text="Copié dans le presse-papier !")

    def save_result(self):
        """Sauvegarde le résultat dans un fichier (.txt ou .docx)."""
        if not self.current_file:
            return

        default_name = self.current_file.stem + "_transcription"

        # Construire la liste des types de fichiers
        filetypes = [("Fichiers texte", "*.txt")]
        if DOCX_AVAILABLE:
            filetypes.insert(0, ("Document Word", "*.docx"))
        filetypes.append(("Tous les fichiers", "*.*"))

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx" if DOCX_AVAILABLE else ".txt",
            initialfile=default_name,
            filetypes=filetypes
        )

        if file_path:
            content = self.result_text.get(1.0, tk.END).strip()
            path = Path(file_path)

            if path.suffix.lower() == ".docx":
                if not DOCX_AVAILABLE:
                    messagebox.showerror(
                        "Erreur",
                        "python-docx n'est pas installé.\n"
                        "Installez-le avec: pip install python-docx"
                    )
                    return
                self._save_as_docx(path, content)
            else:
                path.write_text(content, encoding="utf-8")

            self.status_label.configure(text=f"Sauvegardé: {file_path}")

    def _save_as_docx(self, path: Path, content: str):
        """Exporte la transcription en document Word."""
        doc = Document()

        # Titre du document
        title = doc.add_heading(f"Transcription: {self.current_file.name}", level=1)

        # Ajouter le contenu
        for paragraph_text in content.split("\n\n"):
            if paragraph_text.strip():
                p = doc.add_paragraph(paragraph_text.strip())
                p.style.font.size = Pt(11)

        doc.save(str(path))

    def run(self):
        """Lance l'application."""
        self.root.mainloop()


def main():
    app = TranscriptionApp()
    app.run()


if __name__ == "__main__":
    main()
