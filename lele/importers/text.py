"""Importer pour les fichiers texte, PDF et Word."""

from pathlib import Path

from .base import BaseImporter, ImportResult
from ..models.source import Source, SourceType


class TextImporter(BaseImporter):
    """Importe les fichiers texte, PDF et documents Word."""

    def import_file(
        self,
        file_path: Path,
        project_files_path: Path,
        **options,
    ) -> ImportResult:
        """Importe un fichier texte."""
        if isinstance(file_path, str):
            file_path = Path(file_path)

        valid, error = self.validate_file(file_path)
        if not valid:
            return ImportResult(success=False, error=error)

        self.report_progress(0.1, "Lecture du fichier...")

        ext = file_path.suffix.lower()
        warnings = []

        try:
            # Extraire le contenu selon le format
            if ext == ".pdf":
                content, meta_extra = self._extract_pdf(file_path)
                source_type = SourceType.PDF
            elif ext in (".doc", ".docx", ".odt"):
                content, meta_extra = self._extract_word(file_path)
                source_type = SourceType.WORD
            else:
                content, meta_extra = self._extract_text(file_path)
                source_type = SourceType.TEXT

            self.report_progress(0.5, "Copie du fichier...")

            # Copier dans le projet
            dest_path = self.copy_to_project(file_path, project_files_path)

            self.report_progress(0.8, "Création de la source...")

            # Créer la source
            metadata = self.get_file_metadata(file_path)
            metadata.update(meta_extra)

            source = Source(
                name=file_path.stem,
                type=source_type,
                file_path=str(dest_path),
                content=content,
                metadata=metadata,
            )

            self.report_progress(1.0, "Import terminé")

            return ImportResult(
                success=True,
                source=source,
                warnings=warnings,
                metadata=metadata,
            )

        except Exception as e:
            return ImportResult(success=False, error=str(e))

    def _extract_text(self, file_path: Path) -> tuple[str, dict]:
        """Extrait le texte d'un fichier texte simple."""
        encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

        for encoding in encodings:
            try:
                content = file_path.read_text(encoding=encoding)
                return content, {"encoding": encoding}
            except UnicodeDecodeError:
                continue

        # Fallback avec remplacement des erreurs
        content = file_path.read_text(encoding="utf-8", errors="replace")
        return content, {"encoding": "utf-8", "encoding_errors": True}

    def _extract_pdf(self, file_path: Path) -> tuple[str, dict]:
        """Extrait le texte d'un PDF."""
        try:
            import pypdf

            reader = pypdf.PdfReader(str(file_path))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)

            content = "\n\n".join(pages)
            metadata = {
                "page_count": len(reader.pages),
                "pdf_metadata": dict(reader.metadata) if reader.metadata else {},
            }
            return content, metadata

        except ImportError:
            # Fallback avec pdfplumber
            try:
                import pdfplumber

                pages = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            pages.append(text)

                content = "\n\n".join(pages)
                return content, {"page_count": len(pdf.pages)}

            except ImportError:
                raise ImportError(
                    "Installez pypdf ou pdfplumber pour lire les PDF: "
                    "pip install pypdf ou pip install pdfplumber"
                )

    def _extract_word(self, file_path: Path) -> tuple[str, dict]:
        """Extrait le texte d'un document Word."""
        ext = file_path.suffix.lower()

        if ext == ".docx":
            return self._extract_docx(file_path)
        elif ext == ".doc":
            return self._extract_doc(file_path)
        elif ext == ".odt":
            return self._extract_odt(file_path)
        else:
            raise ValueError(f"Format Word non supporté: {ext}")

    def _extract_docx(self, file_path: Path) -> tuple[str, dict]:
        """Extrait le texte d'un fichier .docx."""
        try:
            import docx

            doc = docx.Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs]
            content = "\n\n".join(paragraphs)

            # Extraire aussi les tableaux
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    tables_text.append("\t".join(cells))

            if tables_text:
                content += "\n\n--- Tableaux ---\n" + "\n".join(tables_text)

            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            }
            return content, metadata

        except ImportError:
            raise ImportError(
                "Installez python-docx pour lire les fichiers Word: "
                "pip install python-docx"
            )

    def _extract_doc(self, file_path: Path) -> tuple[str, dict]:
        """Extrait le texte d'un fichier .doc (ancien format)."""
        try:
            import subprocess

            # Essayer avec antiword
            result = subprocess.run(
                ["antiword", str(file_path)],
                capture_output=True,
                text=True,
            )
            if result.returncode == 0:
                return result.stdout, {"method": "antiword"}

        except FileNotFoundError:
            pass

        # Fallback: lire comme binaire et extraire le texte visible
        try:
            content = file_path.read_bytes()
            # Extraction basique des chaînes ASCII
            import re

            text_parts = re.findall(b"[\x20-\x7e]{4,}", content)
            text = "\n".join(p.decode("ascii", errors="ignore") for p in text_parts)
            return text, {"method": "basic_extraction", "warning": "Extraction basique"}

        except Exception as e:
            raise ValueError(
                f"Impossible de lire le fichier .doc. Installez antiword ou "
                f"convertissez en .docx: {e}"
            )

    def _extract_odt(self, file_path: Path) -> tuple[str, dict]:
        """Extrait le texte d'un fichier OpenDocument."""
        try:
            from odf import text as odf_text
            from odf.opendocument import load

            doc = load(str(file_path))
            paragraphs = doc.getElementsByType(odf_text.P)
            content = "\n\n".join(
                "".join(str(node) for node in p.childNodes)
                for p in paragraphs
            )
            return content, {"paragraph_count": len(paragraphs)}

        except ImportError:
            # Fallback: extraire directement du XML
            import zipfile
            import xml.etree.ElementTree as ET

            with zipfile.ZipFile(file_path, "r") as z:
                content_xml = z.read("content.xml")

            root = ET.fromstring(content_xml)
            # Namespace ODT
            ns = {"text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0"}

            paragraphs = []
            for p in root.iter():
                if p.tag.endswith("}p") or p.tag.endswith("}h"):
                    text = "".join(p.itertext())
                    if text.strip():
                        paragraphs.append(text)

            return "\n\n".join(paragraphs), {"method": "xml_extraction"}
