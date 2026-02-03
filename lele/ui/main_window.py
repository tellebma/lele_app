"""Fenêtre principale de l'application Lele."""

import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from pathlib import Path
from typing import Optional
import traceback

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    DND_AVAILABLE = True
except ImportError:
    DND_AVAILABLE = False

from .. import get_logger, __version__
from ..models.project import Project
from ..models.source import Source, SourceType
from ..models.node import Node
from ..models.coding import CodeReference
from ..importers import get_importer
from ..utils.settings import get_settings_manager
from .dialogs import (
    TranscriptionSettingsDialog,
    ImportProgressDialog,
    AutoCodingConfigDialog,
    AutoCodingPreviewDialog,
    AutoCodingProgressDialog,
    LLMSettingsDialog,
)

# Logger pour ce module
logger = get_logger("ui.main_window")


class MainWindow:
    """Fenêtre principale de l'application QDA."""

    def __init__(self):
        # Créer la fenêtre principale
        if DND_AVAILABLE:
            self.root = TkinterDnD.Tk()
        else:
            self.root = tk.Tk()

        self.root.title("Lele - Analyse Qualitative")
        self.root.geometry("1400x900")
        self.root.minsize(1000, 600)

        # État de l'application
        self.project: Optional[Project] = None
        self.current_source: Optional[Source] = None
        self.selected_node: Optional[Node] = None

        # Mode édition (lecture seule par défaut)
        self.edit_mode = False
        self.original_content: str | None = None  # Pour détecter les changements

        # Gestionnaire de paramètres
        self.settings_manager = get_settings_manager()

        # Paramètres de transcription (chargés depuis les settings)
        settings = self.settings_manager.settings
        self.whisper_model = settings.whisper_model
        self.whisper_language = settings.whisper_language
        self.transcription_show_timestamps = settings.transcription_show_timestamps

        # Configuration des styles
        self.setup_styles()

        # Créer l'interface
        self.setup_menu()
        self.setup_toolbar()
        self.setup_main_layout()
        self.setup_status_bar()

        # Bindings
        self.setup_bindings()

    def setup_styles(self):
        """Configure les styles ttk."""
        style = ttk.Style()
        style.theme_use("clam")

        # Styles personnalisés
        style.configure("Sidebar.TFrame", background="#f5f5f5")
        style.configure("Toolbar.TFrame", background="#e0e0e0")
        style.configure("Title.TLabel", font=("", 11, "bold"))
        style.configure("Sidebar.TLabel", background="#f5f5f5")
        style.configure(
            "Treeview",
            rowheight=25,
            font=("", 10),
        )
        style.configure("Treeview.Heading", font=("", 10, "bold"))

    def setup_menu(self):
        """Configure la barre de menu."""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        # Menu Fichier
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Fichier", menu=file_menu)
        file_menu.add_command(
            label="Nouveau projet...", command=self.new_project, accelerator="Ctrl+N"
        )
        file_menu.add_command(
            label="Ouvrir projet...", command=self.open_project, accelerator="Ctrl+O"
        )

        # Sous-menu Projets récents
        self.recent_menu = tk.Menu(file_menu, tearoff=0)
        file_menu.add_cascade(label="Projets récents", menu=self.recent_menu)
        self._update_recent_projects_menu()

        file_menu.add_command(
            label="Sauvegarder", command=self.save_project, accelerator="Ctrl+S"
        )
        file_menu.add_separator()
        file_menu.add_command(label="Importer...", command=self.import_files)
        file_menu.add_command(label="Exporter...", command=self.export_project)
        file_menu.add_separator()
        file_menu.add_command(label="Quitter", command=self.quit_app, accelerator="Ctrl+Q")

        # Menu Édition
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Édition", menu=edit_menu)
        edit_menu.add_command(label="Annuler", command=self.undo_text, accelerator="Ctrl+Z")
        edit_menu.add_command(label="Rétablir", command=self.redo_text, accelerator="Ctrl+Y")
        edit_menu.add_separator()
        edit_menu.add_command(label="Rechercher...", command=self.show_search, accelerator="Ctrl+F")

        # Menu Codage
        coding_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Codage", menu=coding_menu)
        coding_menu.add_command(label="Nouveau nœud...", command=self.create_node, accelerator="Ctrl+Shift+N")
        coding_menu.add_command(label="Coder la sélection", command=self.code_selection, accelerator="Ctrl+K")
        coding_menu.add_separator()
        coding_menu.add_command(
            label="🔮 Détection automatique de nœuds...",
            command=self.auto_detect_nodes
        )

        # Menu Analyse
        analysis_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Analyse", menu=analysis_menu)
        analysis_menu.add_command(label="Nuage de mots", command=self.show_wordcloud)
        analysis_menu.add_command(label="Carte mentale", command=self.show_mindmap)
        analysis_menu.add_command(label="Sociogramme", command=self.show_sociogram)

        # Menu Paramètres
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Paramètres", menu=settings_menu)
        settings_menu.add_command(
            label="Transcription audio/vidéo...",
            command=self.show_transcription_settings,
        )
        settings_menu.add_command(
            label="🔮 IA / LLM local...",
            command=self.show_llm_settings,
        )

        # Menu Aide
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Aide", menu=help_menu)
        help_menu.add_command(label="📖 Guide d'utilisation", command=self.show_help, accelerator="F1")
        help_menu.add_command(label="⌨️ Raccourcis clavier", command=self.show_shortcuts_help)
        help_menu.add_separator()
        help_menu.add_command(label="À propos de Lele", command=self.show_about)

    def setup_toolbar(self):
        """Configure la barre d'outils."""
        toolbar = ttk.Frame(self.root, style="Toolbar.TFrame", padding="5")
        toolbar.pack(fill=tk.X, side=tk.TOP)

        # Boutons de la barre d'outils
        ttk.Button(toolbar, text="📁 Nouveau", command=self.new_project).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="📂 Ouvrir", command=self.open_project).pack(side=tk.LEFT, padx=2)
        ttk.Button(toolbar, text="💾 Sauver", command=self.save_project).pack(side=tk.LEFT, padx=2)

        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)

        ttk.Button(toolbar, text="📥 Importer", command=self.import_files).pack(side=tk.LEFT, padx=2)

        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)

        ttk.Button(toolbar, text="🏷️ Coder", command=self.code_selection).pack(side=tk.LEFT, padx=2)

        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)

        ttk.Button(toolbar, text="🔍 Rechercher", command=self.show_search).pack(side=tk.LEFT, padx=2)

        ttk.Separator(toolbar, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=10)

        # Boutons mode édition
        self.edit_btn = ttk.Button(toolbar, text="✏️ Éditer", command=self.toggle_edit_mode)
        self.edit_btn.pack(side=tk.LEFT, padx=2)

        self.save_edit_btn = ttk.Button(toolbar, text="💾 Enregistrer", command=self.save_edit)
        # Caché par défaut (affiché en mode édition)

        self.cancel_edit_btn = ttk.Button(toolbar, text="✖️ Annuler", command=self.cancel_edit)
        # Caché par défaut (affiché en mode édition)

        # Recherche rapide
        ttk.Label(toolbar, text="  ").pack(side=tk.LEFT)
        self.quick_search_var = tk.StringVar()
        quick_search = ttk.Entry(toolbar, textvariable=self.quick_search_var, width=30)
        quick_search.pack(side=tk.LEFT, padx=2)
        quick_search.bind("<Return>", lambda e: self.quick_search())

    def setup_main_layout(self):
        """Configure la disposition principale."""
        # PanedWindow principal (horizontal)
        self.main_paned = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        self.main_paned.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Panneau gauche (navigation)
        self.left_panel = ttk.Frame(self.main_paned, width=250)
        self.main_paned.add(self.left_panel, weight=0)

        # Notebook pour la navigation
        self.nav_notebook = ttk.Notebook(self.left_panel)
        self.nav_notebook.pack(fill=tk.BOTH, expand=True)

        # Onglet Sources
        self.sources_frame = ttk.Frame(self.nav_notebook)
        self.nav_notebook.add(self.sources_frame, text="Sources")
        self.setup_sources_panel()

        # Onglet Nœuds
        self.nodes_frame = ttk.Frame(self.nav_notebook)
        self.nav_notebook.add(self.nodes_frame, text="Nœuds")
        self.setup_nodes_panel()


        # Panneau central (contenu)
        self.center_paned = ttk.PanedWindow(self.main_paned, orient=tk.VERTICAL)
        self.main_paned.add(self.center_paned, weight=1)

        # Zone de contenu principale
        self.content_frame = ttk.Frame(self.center_paned)
        self.center_paned.add(self.content_frame, weight=1)
        self.setup_content_area()

        # Zone d'analyse en bas
        self.analysis_frame = ttk.Frame(self.center_paned, height=200)
        self.center_paned.add(self.analysis_frame, weight=0)
        self.setup_analysis_area()

        # Panneau droit (propriétés/codage)
        self.right_panel = ttk.Frame(self.main_paned, width=300)
        self.main_paned.add(self.right_panel, weight=0)
        self.setup_right_panel()

    def setup_sources_panel(self):
        """Configure le panneau des sources."""
        # Boutons
        btn_frame = ttk.Frame(self.sources_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(btn_frame, text="+", width=3, command=self.import_files).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="-", width=3, command=self.delete_source).pack(side=tk.LEFT, padx=2)

        # Filtre par type
        self.source_type_filter = ttk.Combobox(
            btn_frame,
            values=["Tous", "Texte", "Audio", "Vidéo", "Image", "Tableur", "PDF"],
            state="readonly",
            width=10,
        )
        self.source_type_filter.set("Tous")
        self.source_type_filter.pack(side=tk.RIGHT)
        self.source_type_filter.bind("<<ComboboxSelected>>", lambda e: self.refresh_sources())

        # Arbre des sources
        self.sources_tree = ttk.Treeview(
            self.sources_frame,
            columns=("type", "refs"),
            show="tree headings",
            selectmode="browse",
        )
        self.sources_tree.heading("#0", text="Nom")
        self.sources_tree.heading("type", text="Type")
        self.sources_tree.heading("refs", text="Réf.")
        self.sources_tree.column("#0", width=150)
        self.sources_tree.column("type", width=60)
        self.sources_tree.column("refs", width=40)

        sources_scroll = ttk.Scrollbar(self.sources_frame, orient=tk.VERTICAL, command=self.sources_tree.yview)
        self.sources_tree.configure(yscrollcommand=sources_scroll.set)

        self.sources_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(5, 0), pady=5)
        sources_scroll.pack(side=tk.RIGHT, fill=tk.Y, pady=5)

        self.sources_tree.bind("<<TreeviewSelect>>", self.on_source_select)
        self.sources_tree.bind("<Double-1>", self.on_source_double_click)
        self.sources_tree.bind("<Button-3>", self._show_source_context_menu)

        # Menu contextuel pour les sources
        self.source_context_menu = tk.Menu(self.sources_tree, tearoff=0)

        # Drop zone pour le drag & drop
        if DND_AVAILABLE:
            self.sources_tree.drop_target_register(DND_FILES)
            self.sources_tree.dnd_bind("<<Drop>>", self.on_files_drop)

    def setup_nodes_panel(self):
        """Configure le panneau des nœuds."""
        # Boutons
        btn_frame = ttk.Frame(self.nodes_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(btn_frame, text="+", width=3, command=self.create_node).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="-", width=3, command=self.delete_node).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="📁", width=3, command=self.create_node_folder).pack(side=tk.LEFT, padx=2)

        # Bouton de détection automatique
        ttk.Button(
            btn_frame,
            text="🔮 Auto",
            width=8,
            command=self.auto_detect_nodes,
        ).pack(side=tk.RIGHT)

        # Arbre des nœuds
        self.nodes_tree = ttk.Treeview(
            self.nodes_frame,
            columns=("refs",),
            show="tree headings",
            selectmode="browse",
        )
        self.nodes_tree.heading("#0", text="Nœud")
        self.nodes_tree.heading("refs", text="Réf.")
        self.nodes_tree.column("#0", width=180)
        self.nodes_tree.column("refs", width=40)

        nodes_scroll = ttk.Scrollbar(self.nodes_frame, orient=tk.VERTICAL, command=self.nodes_tree.yview)
        self.nodes_tree.configure(yscrollcommand=nodes_scroll.set)

        self.nodes_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(5, 0), pady=5)
        nodes_scroll.pack(side=tk.RIGHT, fill=tk.Y, pady=5)

        self.nodes_tree.bind("<<TreeviewSelect>>", self.on_node_select)
        self.nodes_tree.bind("<Double-1>", self._on_node_double_click)
        self.nodes_tree.bind("<Button-3>", self._show_node_context_menu)

        # Menu contextuel sur les nœuds
        self.node_context_menu = tk.Menu(self.nodes_tree, tearoff=0)

    def setup_content_area(self):
        """Configure la zone de contenu principale."""
        # Notebook pour différentes vues
        self.content_notebook = ttk.Notebook(self.content_frame)
        self.content_notebook.pack(fill=tk.BOTH, expand=True)

        # Vue texte
        self.text_frame = ttk.Frame(self.content_notebook)
        self.content_notebook.add(self.text_frame, text="Document")

        # Zone de texte avec numéros de ligne
        text_container = ttk.Frame(self.text_frame)
        text_container.pack(fill=tk.BOTH, expand=True)

        self.line_numbers = tk.Text(
            text_container,
            width=4,
            padx=4,
            pady=10,  # Même padding que content_text pour alignement
            takefocus=0,
            border=0,
            background="#f0f0f0",
            state="disabled",
            font=("Consolas", 11),
        )
        self.line_numbers.pack(side=tk.LEFT, fill=tk.Y)

        self.content_text = tk.Text(
            text_container,
            wrap=tk.WORD,
            font=("Consolas", 11),
            padx=10,
            pady=10,
            undo=True,
            maxundo=-1,  # Historique illimité
        )
        self.content_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Scrollbar synchronisée avec le texte ET les numéros de ligne
        self.text_scroll = ttk.Scrollbar(text_container, orient=tk.VERTICAL)
        self.text_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        # Fonction de synchronisation du scroll
        def on_text_scroll(*args):
            """Synchronise le scroll entre le texte et les numéros de ligne."""
            self.line_numbers.yview_moveto(args[0])
            self.text_scroll.set(*args)

        def on_scrollbar_scroll(*args):
            """Applique le scroll aux deux widgets de texte."""
            self.content_text.yview(*args)
            self.line_numbers.yview(*args)

        self.content_text.configure(yscrollcommand=on_text_scroll)
        self.text_scroll.configure(command=on_scrollbar_scroll)

        # Synchroniser aussi le scroll à la molette sur les numéros de ligne
        def on_mousewheel_line_numbers(event):
            self.content_text.yview_scroll(int(-1 * (event.delta / 120)), "units")
            return "break"

        self.line_numbers.bind("<MouseWheel>", on_mousewheel_line_numbers)

        # Mettre à jour les numéros de ligne lors du redimensionnement
        # (le word-wrap peut changer le nombre de lignes visuelles)
        def on_text_configure(event):
            # Utiliser after pour éviter les appels multiples rapides
            if hasattr(self, "_line_numbers_update_pending"):
                self.root.after_cancel(self._line_numbers_update_pending)
            self._line_numbers_update_pending = self.root.after(100, self.update_line_numbers)

        self.content_text.bind("<Configure>", on_text_configure)

        # Tags pour le surlignage des codes
        self.content_text.tag_configure("highlight", background="#fff3cd")
        self.content_text.tag_configure("selection", background="#cce5ff")

        # Menu contextuel sur le texte
        self.text_context_menu = tk.Menu(self.content_text, tearoff=0)
        self.content_text.bind("<Button-3>", self._show_text_context_menu)

        # Vue média (pour audio/vidéo/images)
        self.media_frame = ttk.Frame(self.content_notebook)
        self.content_notebook.add(self.media_frame, text="Média")

        self.media_label = ttk.Label(
            self.media_frame,
            text="Sélectionnez un fichier média",
            anchor=tk.CENTER,
        )
        self.media_label.pack(expand=True)

    def setup_analysis_area(self):
        """Configure la zone d'analyse en bas."""
        # Notebook pour les résultats d'analyse
        self.analysis_notebook = ttk.Notebook(self.analysis_frame)
        self.analysis_notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Résultats de recherche
        self.search_frame = ttk.Frame(self.analysis_notebook)
        self.analysis_notebook.add(self.search_frame, text="Recherche")

        self.search_results = ttk.Treeview(
            self.search_frame,
            columns=("type", "snippet"),
            show="headings",
        )
        self.search_results.heading("type", text="Type")
        self.search_results.heading("snippet", text="Extrait")
        self.search_results.column("type", width=80)
        self.search_results.column("snippet", width=500)
        self.search_results.pack(fill=tk.BOTH, expand=True)

        # Références de codage
        self.refs_frame = ttk.Frame(self.analysis_notebook)
        self.analysis_notebook.add(self.refs_frame, text="Références")

        self.refs_tree = ttk.Treeview(
            self.refs_frame,
            columns=("source", "line", "content"),
            show="headings",
        )
        self.refs_tree.heading("source", text="Source")
        self.refs_tree.heading("line", text="Ligne")
        self.refs_tree.heading("content", text="Contenu")
        self.refs_tree.column("source", width=150)
        self.refs_tree.column("line", width=60, anchor=tk.CENTER)
        self.refs_tree.column("content", width=400)
        self.refs_tree.pack(fill=tk.BOTH, expand=True)

        # Bindings pour navigation et suppression
        self.refs_tree.bind("<Double-1>", self._on_ref_double_click)
        self.refs_tree.bind("<Button-3>", self._show_refs_context_menu)
        self.refs_tree.bind("<Delete>", self._delete_selected_ref)

        # Menu contextuel pour les références
        self.refs_context_menu = tk.Menu(self.refs_tree, tearoff=0)

    def setup_right_panel(self):
        """Configure le panneau droit (codages du document)."""
        # Titre
        ttk.Label(self.right_panel, text="Codages du document", style="Title.TLabel").pack(
            anchor=tk.W, padx=10, pady=(10, 5)
        )

        # Arbre des codages
        self.doc_codes_tree = ttk.Treeview(
            self.right_panel,
            columns=("pos",),
            show="tree headings",
            height=10,
        )
        self.doc_codes_tree.heading("#0", text="Nœud")
        self.doc_codes_tree.heading("pos", text="Position")
        self.doc_codes_tree.column("#0", width=150)
        self.doc_codes_tree.column("pos", width=80)
        self.doc_codes_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        # Bindings pour navigation et suppression
        self.doc_codes_tree.bind("<Double-1>", self._on_doc_code_double_click)
        self.doc_codes_tree.bind("<Button-3>", self._show_doc_codes_context_menu)
        self.doc_codes_tree.bind("<Delete>", self._delete_selected_doc_code)

        # Menu contextuel pour les codages
        self.doc_codes_context_menu = tk.Menu(self.doc_codes_tree, tearoff=0)

    def setup_status_bar(self):
        """Configure la barre de statut."""
        self.status_bar = ttk.Frame(self.root)
        self.status_bar.pack(fill=tk.X, side=tk.BOTTOM)

        self.status_label = ttk.Label(self.status_bar, text="Prêt")
        self.status_label.pack(side=tk.LEFT, padx=10, pady=5)

        # Indicateur de mode édition
        self.mode_label = ttk.Label(
            self.status_bar,
            text="🔒 Lecture seule",
            foreground="#666666",
            font=("", 9),
        )
        self.mode_label.pack(side=tk.LEFT, padx=20, pady=5)

        self.project_label = ttk.Label(self.status_bar, text="Aucun projet")
        self.project_label.pack(side=tk.RIGHT, padx=10, pady=5)

    def setup_bindings(self):
        """Configure les raccourcis clavier."""
        self.root.bind("<Control-n>", lambda e: self.new_project())
        self.root.bind("<Control-o>", lambda e: self.open_project())
        self.root.bind("<Control-s>", lambda e: self.save_project())
        self.root.bind("<Control-f>", lambda e: self.show_search())
        self.root.bind("<Control-k>", lambda e: self.code_selection())
        self.root.bind("<Control-q>", lambda e: self.quit_app())

        # Édition (undo/redo)
        self.root.bind("<Control-z>", lambda e: self.undo_text())
        self.root.bind("<Control-y>", lambda e: self.redo_text())

        # Raccourcis supplémentaires pour les nœuds
        self.root.bind("<Control-Shift-N>", lambda e: self.create_node())
        self.root.bind("<Control-Shift-K>", lambda e: self._quick_code_from_selection())
        self.root.bind("<F2>", lambda e: self._rename_node() if self.selected_node else None)
        self.root.bind("<Delete>", lambda e: self._delete_if_node_focused())

        # Mode édition
        self.root.bind("<Control-e>", lambda e: self.toggle_edit_mode())
        self.root.bind("<Escape>", lambda e: self.cancel_edit() if self.edit_mode else None)

        # Aide
        self.root.bind("<F1>", lambda e: self.show_help())

    # --- Projets récents ---

    def _update_recent_projects_menu(self):
        """Met à jour le menu des projets récents."""
        # Vider le menu actuel
        self.recent_menu.delete(0, tk.END)

        # Nettoyer les projets qui n'existent plus
        self.settings_manager.clean_nonexistent_projects()

        # Récupérer les projets récents
        recent_projects = self.settings_manager.get_recent_projects()

        if not recent_projects:
            self.recent_menu.add_command(
                label="(Aucun projet récent)",
                state=tk.DISABLED,
            )
        else:
            for i, project in enumerate(recent_projects):
                path = project["path"]
                name = project["name"]

                # Raccourci clavier pour les 9 premiers
                accelerator = f"Ctrl+{i + 1}" if i < 9 else ""

                self.recent_menu.add_command(
                    label=f"{name}",
                    command=lambda p=path: self._open_recent_project(p),
                    accelerator=accelerator,
                )

            # Séparateur et option pour effacer
            self.recent_menu.add_separator()
            self.recent_menu.add_command(
                label="Effacer la liste",
                command=self._clear_recent_projects,
            )

    def _open_recent_project(self, path: str):
        """Ouvre un projet récent."""
        project_path = Path(path)

        if not project_path.exists():
            messagebox.showerror(
                "Erreur",
                f"Le projet n'existe plus:\n{path}",
            )
            self.settings_manager.remove_recent_project(path)
            self._update_recent_projects_menu()
            return

        try:
            self.project = Project.open(project_path)
            self.refresh_all()
            self.update_status(f"Projet ouvert: {self.project.name}")
            self.project_label.configure(text=self.project.name)

            # Mettre à jour les projets récents (le met en tête)
            self.settings_manager.add_recent_project(path)
            self._update_recent_projects_menu()

        except Exception as e:
            messagebox.showerror("Erreur", f"Impossible d'ouvrir le projet: {e}")
            logger.error(f"Erreur ouverture projet récent: {e}")

    def _clear_recent_projects(self):
        """Efface la liste des projets récents."""
        if messagebox.askyesno(
            "Confirmer",
            "Effacer la liste des projets récents?",
        ):
            self.settings_manager.clear_recent_projects()
            self._update_recent_projects_menu()

    # --- Actions du menu ---

    def new_project(self):
        """Crée un nouveau projet."""
        dialog = tk.Toplevel(self.root)
        dialog.title("Nouveau projet")
        dialog.geometry("400x200")
        dialog.transient(self.root)
        dialog.grab_set()

        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(main_frame, text="Nom du projet:").pack(anchor=tk.W)
        name_var = tk.StringVar()
        ttk.Entry(main_frame, textvariable=name_var, width=40).pack(pady=5, fill=tk.X)

        ttk.Label(main_frame, text="Emplacement:").pack(pady=(10, 0), anchor=tk.W)
        path_var = tk.StringVar()
        path_frame = ttk.Frame(main_frame)
        path_frame.pack(pady=5, fill=tk.X)
        ttk.Entry(path_frame, textvariable=path_var, width=30).pack(side=tk.LEFT, expand=True, fill=tk.X)
        ttk.Button(
            path_frame,
            text="...",
            width=3,
            command=lambda: path_var.set(filedialog.askdirectory()),
        ).pack(side=tk.LEFT, padx=(5, 0))

        def create():
            name = name_var.get().strip()
            path = path_var.get().strip()
            if not name or not path:
                messagebox.showerror("Erreur", "Veuillez remplir tous les champs")
                return

            project_path = Path(path) / name
            try:
                self.project = Project(name=name, path=project_path)
                self.project.create()
                self.refresh_all()
                self.update_status(f"Projet créé: {name}")
                self.project_label.configure(text=name)

                # Ajouter aux projets récents
                self.settings_manager.add_recent_project(project_path)
                self._update_recent_projects_menu()

                dialog.destroy()
            except Exception as e:
                messagebox.showerror("Erreur", f"Erreur lors de la création: {e}")

        ttk.Button(main_frame, text="Créer", command=create).pack(pady=(15, 0))

    def open_project(self):
        """Ouvre un projet existant."""
        path = filedialog.askdirectory(title="Sélectionner le dossier du projet")
        if path:
            try:
                self.project = Project.open(Path(path))
                self.refresh_all()
                self.update_status(f"Projet ouvert: {self.project.name}")
                self.project_label.configure(text=self.project.name)

                # Ajouter aux projets récents
                self.settings_manager.add_recent_project(path)
                self._update_recent_projects_menu()

            except Exception as e:
                messagebox.showerror("Erreur", f"Impossible d'ouvrir le projet: {e}")

    def save_project(self):
        """Sauvegarde le projet."""
        if self.project:
            self.project.save()
            self.update_status("Projet sauvegardé")

    def import_files(self):
        """Importe des fichiers."""
        if not self.project:
            messagebox.showwarning("Attention", "Veuillez d'abord créer ou ouvrir un projet")
            return

        filetypes = [
            ("Tous les fichiers supportés", "*.txt *.pdf *.docx *.mp3 *.wav *.mp4 *.jpg *.png *.xlsx *.csv"),
            ("Documents texte", "*.txt *.md *.rtf"),
            ("PDF", "*.pdf"),
            ("Word", "*.doc *.docx"),
            ("Audio", "*.mp3 *.wav *.m4a *.flac *.ogg"),
            ("Vidéo", "*.mp4 *.avi *.mov *.mkv"),
            ("Images", "*.jpg *.jpeg *.png *.gif *.bmp"),
            ("Tableurs", "*.xlsx *.xls *.csv"),
            ("Tous les fichiers", "*.*"),
        ]

        files = filedialog.askopenfilenames(filetypes=filetypes)
        if files:
            self.import_files_list(files)

    def import_files_list(self, files):
        """Importe une liste de fichiers."""
        # Vérifier si des fichiers audio/vidéo sont présents
        audio_video_extensions = {
            ".mp3", ".wav", ".m4a", ".flac", ".ogg", ".webm",
            ".mp4", ".avi", ".mov", ".mkv", ".wmv",
        }

        has_audio_video = any(
            Path(f).suffix.lower() in audio_video_extensions for f in files
        )

        # Afficher le dialogue de paramètres si audio/vidéo détecté
        transcribe = True
        whisper_model = self.whisper_model
        whisper_language = self.whisper_language
        show_timestamps = self.transcription_show_timestamps

        if has_audio_video:
            dialog = TranscriptionSettingsDialog(
                self.root,
                current_model=self.whisper_model,
                current_language=self.whisper_language,
                show_transcribe_option=True,
                current_show_timestamps=self.transcription_show_timestamps,
            )
            self.root.wait_window(dialog)

            if dialog.cancelled:
                return

            transcribe = dialog.result_transcribe
            whisper_model = dialog.result_model
            whisper_language = dialog.result_language
            show_timestamps = dialog.result_show_timestamps

            # Sauvegarder les préférences
            self.whisper_model = whisper_model
            self.whisper_language = whisper_language
            self.transcription_show_timestamps = show_timestamps

            # Persister dans les settings
            self.settings_manager.settings.transcription_show_timestamps = show_timestamps
            self.settings_manager.save()

        # Lancer l'import dans un thread séparé
        self._start_async_import(
            files, transcribe, whisper_model, whisper_language, show_timestamps
        )

    def _start_async_import(
        self,
        files: list,
        transcribe: bool,
        whisper_model: str,
        whisper_language: Optional[str],
        show_timestamps: bool = False,
    ):
        """Lance l'import de fichiers de manière asynchrone."""
        logger.info(f"Démarrage de l'import de {len(files)} fichier(s)")

        # Créer le dialogue de progression
        progress_dialog = ImportProgressDialog(self.root, total_files=len(files))

        def update_progress(progress: float, message: str):
            """Callback pour mettre à jour la progression."""
            if not progress_dialog.cancelled:
                progress_dialog.after(
                    0, lambda: progress_dialog.set_step(progress, message)
                )
                # Logger les messages importants (durée, estimation) dans les détails
                if "Audio:" in message or "Vidéo:" in message or "estimé:" in message:
                    progress_dialog.after(
                        0, lambda m=message: progress_dialog.log(f"  ⏱️ {m}")
                    )

        def do_import():
            sources_to_save = []
            errors = []
            total = len(files)

            for i, file_path in enumerate(files, 1):
                # Vérifier si annulé
                if progress_dialog.cancelled:
                    logger.info("Import annulé par l'utilisateur")
                    break

                try:
                    filename = Path(file_path).name
                    progress_dialog.after(
                        0, lambda f=filename, n=i: progress_dialog.set_file(f, n)
                    )
                    progress_dialog.after(
                        0, lambda f=filename: progress_dialog.log(f"Import: {f}")
                    )

                    logger.info(f"Import du fichier: {file_path}")
                    importer = get_importer(file_path)

                    # Configurer le callback de progression
                    importer.set_progress_callback(update_progress)

                    # Déterminer si c'est un fichier audio/vidéo
                    is_audio_video = Path(file_path).suffix.lower() in {
                        ".mp3", ".wav", ".m4a", ".flac", ".ogg", ".webm",
                        ".mp4", ".avi", ".mov", ".mkv", ".wmv",
                    }

                    # Préparer les options d'import
                    import_options = {}
                    if is_audio_video:
                        import_options["transcribe"] = transcribe
                        import_options["whisper_model"] = whisper_model
                        import_options["language"] = whisper_language
                        import_options["show_timestamps"] = show_timestamps
                        logger.info(
                            f"Options transcription: model={whisper_model}, "
                            f"lang={whisper_language}, transcribe={transcribe}, "
                            f"timestamps={show_timestamps}"
                        )
                        progress_dialog.after(
                            0,
                            lambda: progress_dialog.log(
                                f"  Transcription: modèle={whisper_model}, "
                                f"timestamps={'oui' if show_timestamps else 'non'}"
                            ),
                        )

                    result = importer.import_file(
                        Path(file_path),
                        self.project.files_path,
                        **import_options,
                    )

                    if result.success and result.source:
                        # Collecter la source pour la sauvegarder dans le thread principal
                        sources_to_save.append(result.source)
                        content_len = len(result.source.content or "")
                        logger.info(f"Import réussi: {result.source.name} ({content_len} chars)")
                        progress_dialog.after(
                            0,
                            lambda cl=content_len: progress_dialog.log(
                                f"  ✓ Succès ({cl} caractères)"
                            ),
                        )

                        # Afficher les avertissements
                        if result.warnings:
                            for warning in result.warnings:
                                logger.warning(f"Avertissement: {warning}")
                                progress_dialog.after(
                                    0,
                                    lambda w=warning: progress_dialog.log(f"  ⚠ {w}"),
                                )
                    else:
                        error_msg = f"{filename}: {result.error}"
                        errors.append(error_msg)
                        logger.error(f"Échec de l'import: {error_msg}")
                        progress_dialog.after(
                            0,
                            lambda e=result.error: progress_dialog.log(f"  ✗ Erreur: {e}"),
                        )

                except Exception as e:
                    error_msg = f"{Path(file_path).name}: {e}"
                    errors.append(error_msg)
                    logger.error(f"Exception lors de l'import: {error_msg}")
                    logger.error(traceback.format_exc())
                    progress_dialog.after(
                        0, lambda err=str(e): progress_dialog.log(f"  ✗ Exception: {err}")
                    )

            # Mise à jour de l'interface dans le thread principal
            # Les sources seront sauvegardées dans le thread principal pour éviter
            # les erreurs SQLite "objects created in a thread..."
            self.root.after(
                0, lambda: self._on_import_complete(sources_to_save, errors, progress_dialog)
            )

        thread = threading.Thread(target=do_import, daemon=True)
        thread.start()

    def _on_import_complete(self, sources: list, errors: list, progress_dialog: ImportProgressDialog):
        """Callback appelé quand l'import est terminé."""
        # Sauvegarder les sources dans le thread principal (évite les erreurs SQLite)
        saved_count = 0
        for source in sources:
            try:
                source.save(self.project.db)
                saved_count += 1
                logger.info(f"Source sauvegardée: {source.name}")
                progress_dialog.log(f"💾 Sauvegardé: {source.name}")
            except Exception as e:
                error_msg = f"{source.name}: {e}"
                errors.append(error_msg)
                logger.error(f"Erreur de sauvegarde: {error_msg}")
                logger.error(traceback.format_exc())
                progress_dialog.log(f"✗ Erreur sauvegarde: {source.name}")

        self.refresh_sources()
        self.update_status(f"{saved_count} fichier(s) importé(s)")
        logger.info(f"Import terminé: {saved_count} fichier(s) importé(s)")

        # Mettre à jour le dialogue de progression
        progress_dialog.complete(saved_count, len(errors))

        if errors:
            error_text = "\n".join(errors)
            logger.warning(f"Erreurs d'import:\n{error_text}")

    def on_files_drop(self, event):
        """Gère le drop de fichiers."""
        files = self.root.tk.splitlist(event.data)
        if files and self.project:
            self.import_files_list(files)

    def export_project(self):
        """Exporte le projet au format REFI-QDA (.qdpx)."""
        if not self.project:
            messagebox.showwarning("Export", "Aucun projet ouvert.")
            return

        # Dialogue de sauvegarde
        from pathlib import Path
        filepath = filedialog.asksaveasfilename(
            title="Exporter le projet",
            defaultextension=".qdpx",
            filetypes=[
                ("REFI-QDA Project", "*.qdpx"),
                ("Tous les fichiers", "*.*")
            ],
            initialfile=f"{self.project.name}.qdpx"
        )

        if not filepath:
            return  # Annulé par l'utilisateur

        # Exporter via RefiQdaImporter
        from lele.importers.refi_qda import RefiQdaImporter

        try:
            exporter = RefiQdaImporter()
            success = exporter.export_project(self.project, Path(filepath))

            if success:
                messagebox.showinfo(
                    "Export réussi",
                    f"Le projet a été exporté avec succès vers:\n{filepath}"
                )
            else:
                messagebox.showerror(
                    "Erreur d'export",
                    "L'export a échoué. Vérifiez les permissions du fichier."
                )
        except Exception as e:
            messagebox.showerror(
                "Erreur d'export",
                f"Une erreur est survenue lors de l'export:\n{str(e)}"
            )

    def quit_app(self):
        """Quitte l'application."""
        if self.project:
            self.project.close()
        self.root.quit()

    # --- Édition (undo/redo) ---

    def undo_text(self):
        """Annule la dernière modification dans le texte."""
        try:
            self.content_text.edit_undo()
        except tk.TclError:
            pass  # Rien à annuler

    def redo_text(self):
        """Rétablit la dernière modification annulée."""
        try:
            self.content_text.edit_redo()
        except tk.TclError:
            pass  # Rien à rétablir

    # --- Mode édition ---

    def toggle_edit_mode(self):
        """Bascule entre mode lecture et mode édition."""
        if not self.current_source:
            messagebox.showinfo("Information", "Aucun document ouvert.")
            return

        if self.edit_mode:
            # Vérifier s'il y a des modifications non sauvegardées
            if self._has_unsaved_changes():
                result = messagebox.askyesnocancel(
                    "Modifications non sauvegardées",
                    "Voulez-vous enregistrer les modifications avant de quitter le mode édition?",
                )
                if result is None:  # Cancel
                    return
                elif result:  # Yes - sauvegarder
                    self.save_edit()
                    return
                # No - annuler les modifications
                self.cancel_edit()
                return
            self._set_edit_mode(False)
        else:
            self._set_edit_mode(True)

    def _set_edit_mode(self, enabled: bool):
        """Active ou désactive le mode édition avec tous les changements visuels."""
        self.edit_mode = enabled

        if enabled:
            # Activer le mode édition
            self.original_content = self.content_text.get("1.0", tk.END)
            self.content_text.configure(
                state=tk.NORMAL,
                background="#ffffff",
                relief=tk.SOLID,
                borderwidth=2,
            )

            # Mettre à jour les boutons
            self.edit_btn.configure(text="✏️ Édition", state=tk.DISABLED)
            self.save_edit_btn.pack(side=tk.LEFT, padx=2, after=self.edit_btn)
            self.cancel_edit_btn.pack(side=tk.LEFT, padx=2, after=self.save_edit_btn)

            # Mettre à jour l'indicateur de statut
            self.mode_label.configure(
                text="✏️ Mode édition",
                foreground="#0066cc",
                font=("", 9, "bold"),
            )

            self.update_status("Mode édition activé - Ctrl+S pour sauvegarder, Échap pour annuler")

        else:
            # Désactiver le mode édition
            self.original_content = None
            self.content_text.configure(
                state=tk.DISABLED,
                background="#f8f8f8",
                relief=tk.FLAT,
                borderwidth=0,
            )

            # Mettre à jour les boutons
            self.edit_btn.configure(text="✏️ Éditer", state=tk.NORMAL)
            self.save_edit_btn.pack_forget()
            self.cancel_edit_btn.pack_forget()

            # Mettre à jour l'indicateur de statut
            self.mode_label.configure(
                text="🔒 Lecture seule",
                foreground="#666666",
                font=("", 9),
            )

            self.update_status("Mode lecture seule")

    def _has_unsaved_changes(self) -> bool:
        """Vérifie s'il y a des modifications non sauvegardées."""
        if not self.edit_mode or self.original_content is None:
            return False
        current_content = self.content_text.get("1.0", tk.END)
        return current_content != self.original_content

    def save_edit(self):
        """Enregistre les modifications du document."""
        if not self.edit_mode or not self.current_source:
            return

        # Récupérer le contenu modifié
        new_content = self.content_text.get("1.0", tk.END).rstrip("\n")

        # Mettre à jour la source
        self.current_source.content = new_content
        self.current_source.save(self.project.db)

        # Réinitialiser le contenu original
        self.original_content = self.content_text.get("1.0", tk.END)

        self.update_status("Modifications enregistrées")

        # Quitter le mode édition
        self._set_edit_mode(False)

        # Rafraîchir l'affichage
        self.display_source()

    def cancel_edit(self):
        """Annule les modifications et restaure le contenu original."""
        if not self.edit_mode:
            return

        if self._has_unsaved_changes():
            if not messagebox.askyesno(
                "Annuler les modifications",
                "Voulez-vous vraiment annuler toutes les modifications?"
            ):
                return

        # Restaurer le contenu original
        if self.original_content is not None:
            self.content_text.configure(state=tk.NORMAL)
            self.content_text.delete("1.0", tk.END)
            self.content_text.insert("1.0", self.original_content.rstrip("\n"))

        # Quitter le mode édition
        self._set_edit_mode(False)

        self.update_status("Modifications annulées")

    # --- Actions de codage ---

    def create_node(self, initial_name: str = "", parent_node: Node | None = None, code_selection: bool = False):
        """Crée un nouveau nœud avec dialogue amélioré.

        Args:
            initial_name: Nom pré-rempli (ex: depuis sélection de texte)
            parent_node: Nœud parent (None = racine, sinon enfant)
            code_selection: Si True, code automatiquement la sélection après création
        """
        if not self.project:
            return

        dialog = tk.Toplevel(self.root)
        dialog.title("Nouveau nœud")
        dialog.geometry("450x420")
        dialog.transient(self.root)
        dialog.grab_set()

        # Frame principal avec style ttk standard
        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Nom
        ttk.Label(main_frame, text="Nom du nœud:", font=("", 10, "bold")).pack(
            padx=20, pady=(20, 5), anchor=tk.W
        )
        name_var = tk.StringVar(value=initial_name)
        name_entry = ttk.Entry(main_frame, textvariable=name_var, width=45)
        name_entry.pack(padx=20, pady=5)
        name_entry.focus_set()
        name_entry.select_range(0, tk.END)

        # Description
        ttk.Label(main_frame, text="Description (optionnel):").pack(pady=(10, 5), anchor=tk.W)
        desc_var = tk.StringVar()
        ttk.Entry(main_frame, textvariable=desc_var, width=45).pack(pady=5)

        # Couleur
        ttk.Label(main_frame, text="Couleur:", font=("", 10, "bold")).pack(pady=(15, 5), anchor=tk.W)
        color_var = tk.StringVar(value="#3498db")

        color_frame = ttk.Frame(main_frame)
        color_frame.pack(pady=5, fill=tk.X)

        # Palette de couleurs étendue
        colors = [
            "#e74c3c", "#e91e63", "#9b59b6", "#673ab7",
            "#3498db", "#2196f3", "#00bcd4", "#009688",
            "#2ecc71", "#4caf50", "#8bc34a", "#cddc39",
            "#f39c12", "#ff9800", "#ff5722", "#795548",
        ]

        self._color_buttons = []
        for i, color in enumerate(colors):
            btn = tk.Button(
                color_frame,
                bg=color,
                activebackground=color,
                width=2,
                height=1,
                relief=tk.FLAT,
                cursor="hand2",
                command=lambda c=color: self._select_node_color(c, color_var),
            )
            btn.grid(row=i // 8, column=i % 8, padx=2, pady=2)
            self._color_buttons.append((btn, color))

        # Indicateur de couleur sélectionnée
        preview_frame = ttk.Frame(main_frame)
        preview_frame.pack(pady=10, fill=tk.X)

        ttk.Label(preview_frame, text="Aperçu:").pack(side=tk.LEFT)
        self._color_preview = ttk.Label(
            preview_frame,
            text="  ●  Exemple de nœud",
            foreground=color_var.get(),
            font=("", 11),
        )
        self._color_preview.pack(side=tk.LEFT, padx=10)

        def update_preview(*args):
            self._color_preview.configure(foreground=color_var.get())
        color_var.trace_add("write", update_preview)

        # Parent
        if parent_node:
            parent_label = ttk.Label(
                main_frame,
                text=f"📁 Parent: {parent_node.name}",
                foreground="#666",
            )
            parent_label.pack(anchor=tk.W)

        # Boutons
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=20, fill=tk.X)

        def create():
            name = name_var.get().strip()
            if not name:
                messagebox.showwarning("Attention", "Le nom du nœud est requis.", parent=dialog)
                return

            node = Node(
                name=name,
                description=desc_var.get().strip(),
                color=color_var.get(),
                parent_id=parent_node.id if parent_node else None,
            )
            node.save(self.project.db)
            self.refresh_nodes()

            # Si demandé, coder la sélection avec ce nouveau nœud
            if code_selection:
                self.selected_node = node
                self.code_selection()

            dialog.destroy()
            self.update_status(f"Nœud '{name}' créé")

        ttk.Button(btn_frame, text="Annuler", command=dialog.destroy).pack(side=tk.RIGHT, padx=5)
        create_btn = ttk.Button(btn_frame, text="✓ Créer", command=create)
        create_btn.pack(side=tk.RIGHT)

        # Raccourci Entrée pour créer
        dialog.bind("<Return>", lambda e: create())
        dialog.bind("<Escape>", lambda e: dialog.destroy())

    def _select_node_color(self, color: str, color_var: tk.StringVar):
        """Met à jour la couleur sélectionnée."""
        color_var.set(color)

    def create_node_folder(self):
        """Crée un dossier de nœuds (nœud enfant du nœud sélectionné)."""
        parent = self.selected_node if self.selected_node else None
        self.create_node(parent_node=parent)

    # --- Menus contextuels ---

    def _show_text_context_menu(self, event):
        """Affiche le menu contextuel sur le texte."""
        self.text_context_menu.delete(0, tk.END)

        # Détecter si on clique sur une zone codée
        click_index = self.content_text.index(f"@{event.x},{event.y}")
        clicked_ref = self._get_reference_at_position(click_index)

        # Options pour une référence existante
        if clicked_ref and self.project:
            node = Node.get(self.project.db, clicked_ref.node_id)
            node_name = node.name if node else "inconnu"

            self.text_context_menu.add_command(
                label=f"🏷️ Référence: {node_name}",
                state=tk.DISABLED,
            )
            self.text_context_menu.add_command(
                label="🔄 Changer le nœud...",
                command=lambda r=clicked_ref: self._change_reference_node(r),
            )
            self.text_context_menu.add_command(
                label="🗑️ Supprimer cette référence",
                command=lambda r=clicked_ref: self._delete_reference(r),
            )
            self.text_context_menu.add_separator()

        # Vérifier s'il y a une sélection
        has_selection = False
        selected_text = ""
        try:
            selected_text = self.content_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            has_selection = bool(selected_text.strip())
        except tk.TclError:
            pass

        if has_selection and self.project and self.current_source:
            # Options de codage
            nodes = Node.get_all(self.project.db)

            if nodes:
                # Sous-menu "Coder avec..."
                code_menu = tk.Menu(self.text_context_menu, tearoff=0)
                for node in nodes[:15]:  # Limiter à 15 pour éviter un menu trop long
                    code_menu.add_command(
                        label=f"● {node.name}",
                        foreground=node.color,
                        command=lambda n=node: self._code_with_node(n),
                    )

                if len(nodes) > 15:
                    code_menu.add_separator()
                    code_menu.add_command(label="Plus...", command=self._show_all_nodes_for_coding)

                self.text_context_menu.add_cascade(label="🏷️ Coder avec", menu=code_menu)
                self.text_context_menu.add_separator()

            # Créer un nœud depuis la sélection
            short_text = selected_text[:30].replace("\n", " ")
            if len(selected_text) > 30:
                short_text += "..."

            self.text_context_menu.add_command(
                label="➕ Nouveau nœud depuis la sélection...",
                command=lambda: self.create_node(initial_name=short_text, code_selection=True),
            )

            if self.selected_node:
                self.text_context_menu.add_command(
                    label=f"🏷️ Coder avec '{self.selected_node.name}'",
                    command=self.code_selection,
                )

            self.text_context_menu.add_separator()

        # Options standard
        self.text_context_menu.add_command(
            label="📋 Copier",
            command=lambda: self.root.event_generate("<<Copy>>"),
            accelerator="Ctrl+C",
        )

        if has_selection:
            self.text_context_menu.add_command(
                label="🔍 Rechercher dans les sources",
                command=lambda: self._search_selected_text(selected_text),
            )

        # Options de reformatage pour les sources audio/vidéo avec segments
        if self.current_source and self._has_transcription_segments():
            self.text_context_menu.add_separator()
            reformat_menu = tk.Menu(self.text_context_menu, tearoff=0)
            reformat_menu.add_command(
                label="Avec horodatages [00:00 -> 00:05]",
                command=lambda: self._reformat_transcription(show_timestamps=True),
            )
            reformat_menu.add_command(
                label="Sans horodatages (paragraphes)",
                command=lambda: self._reformat_transcription(show_timestamps=False),
            )
            self.text_context_menu.add_cascade(label="🔄 Reformater la transcription", menu=reformat_menu)

        self.text_context_menu.tk_popup(event.x_root, event.y_root)

    def _has_transcription_segments(self) -> bool:
        """Vérifie si la source actuelle a des segments de transcription."""
        if not self.current_source:
            return False
        transcription = self.current_source.metadata.get("transcription", {})
        segments = transcription.get("segments", [])
        return len(segments) > 0

    def _reformat_transcription(self, show_timestamps: bool):
        """
        Reformate la transcription de la source actuelle.

        Cette opération met à jour le contenu de la source avec le nouveau formatage
        (avec ou sans horodatages) tout en préservant les segments originaux.
        """
        if not self.current_source or not self.project:
            return

        transcription = self.current_source.metadata.get("transcription", {})
        segments = transcription.get("segments", [])

        if not segments:
            messagebox.showwarning(
                "Reformatage impossible",
                "Cette source ne contient pas de segments de transcription."
            )
            return

        # Demander confirmation
        format_type = "avec horodatages" if show_timestamps else "sans horodatages (paragraphes)"
        confirm = messagebox.askyesno(
            "Reformater la transcription",
            f"Voulez-vous reformater le texte {format_type} ?\n\n"
            "Le contenu sera mis à jour mais les segments originaux seront préservés.\n"
            "Note : Les codages existants pourraient être décalés.",
            icon="question"
        )

        if not confirm:
            return

        # Importer l'AudioImporter pour utiliser la méthode de formatage
        from ..importers.audio import AudioImporter

        importer = AudioImporter()
        new_content = importer._format_transcript(segments, show_timestamps)

        # Mettre à jour la source
        self.current_source.content = new_content
        transcription["show_timestamps"] = show_timestamps
        self.current_source.metadata["transcription"] = transcription

        # Sauvegarder
        self.current_source.save(self.project.db)

        # Rafraîchir l'affichage
        self.display_source()

        messagebox.showinfo(
            "Reformatage terminé",
            f"Le texte a été reformaté {format_type}."
        )

    def _show_node_context_menu(self, event):
        """Affiche le menu contextuel sur un nœud."""
        # Sélectionner le nœud sous le curseur
        item = self.nodes_tree.identify_row(event.y)
        if item:
            self.nodes_tree.selection_set(item)
            self.on_node_select(None)

        self.node_context_menu.delete(0, tk.END)

        if self.selected_node:
            self.node_context_menu.add_command(
                label="✏️ Renommer...",
                command=self._rename_node,
            )
            self.node_context_menu.add_command(
                label="🎨 Changer la couleur...",
                command=self._change_node_color,
            )
            self.node_context_menu.add_separator()
            self.node_context_menu.add_command(
                label="➕ Ajouter un sous-nœud...",
                command=lambda: self.create_node(parent_node=self.selected_node),
            )
            self.node_context_menu.add_separator()
            self.node_context_menu.add_command(
                label="🗑️ Supprimer",
                command=self.delete_node,
            )
        else:
            self.node_context_menu.add_command(
                label="➕ Nouveau nœud...",
                command=self.create_node,
            )

        self.node_context_menu.tk_popup(event.x_root, event.y_root)

    def _on_node_double_click(self, event):
        """Double-clic sur un nœud : code la sélection si du texte est sélectionné."""
        if not self.selected_node or not self.current_source:
            return

        try:
            selected_text = self.content_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            if selected_text.strip():
                self.code_selection()
        except tk.TclError:
            # Pas de sélection, ne rien faire
            pass

    def _code_with_node(self, node: Node):
        """Code la sélection avec le nœud spécifié."""
        old_selected = self.selected_node
        self.selected_node = node
        self.code_selection()
        self.selected_node = old_selected

    # ===== Handlers pour les références (refs_tree) =====

    def _on_ref_double_click(self, event):
        """Double-clic sur une référence : naviguer vers la source et position."""
        selection = self.refs_tree.selection()
        if not selection:
            return

        item_id = selection[0]
        ref_id = self.refs_tree.item(item_id, "tags")
        if not ref_id:
            return
        ref_id = ref_id[0]

        # Récupérer la référence
        ref = CodeReference.get(self.project.db, ref_id)
        if not ref:
            return

        # Charger la source correspondante
        source = Source.get(self.project.db, ref.source_id)
        if not source:
            return

        # Afficher la source
        self.current_source = source
        self.display_source()

        # Sélectionner l'item dans sources_tree
        for item in self.sources_tree.get_children():
            if self.sources_tree.item(item, "tags") and self.sources_tree.item(item, "tags")[0] == source.id:
                self.sources_tree.selection_set(item)
                self.sources_tree.see(item)
                break

        # Naviguer vers la position dans le texte
        self._scroll_to_position(ref.start_pos, ref.end_pos)

    def _scroll_to_position(self, start_pos: int, end_pos: int):
        """Scroll le texte vers une position et surligne le passage."""
        # Convertir les positions en indices Tkinter
        start_index = f"1.0+{start_pos}c"
        end_index = f"1.0+{end_pos}c"

        # Supprimer l'ancien surlignage de sélection
        self.content_text.tag_remove("selection", "1.0", tk.END)

        # Surligner le passage
        self.content_text.tag_add("selection", start_index, end_index)

        # Scroll pour rendre visible
        self.content_text.see(start_index)
        self.content_text.mark_set(tk.INSERT, start_index)

    def _show_refs_context_menu(self, event):
        """Affiche le menu contextuel pour les références."""
        # Sélectionner l'item sous le curseur
        item = self.refs_tree.identify_row(event.y)
        if item:
            self.refs_tree.selection_set(item)

        self.refs_context_menu.delete(0, tk.END)

        if item:
            self.refs_context_menu.add_command(
                label="📍 Aller à la source",
                command=lambda: self._on_ref_double_click(None)
            )
            self.refs_context_menu.add_separator()
            self.refs_context_menu.add_command(
                label="🗑️ Supprimer cette référence",
                command=self._delete_selected_ref
            )

        self.refs_context_menu.tk_popup(event.x_root, event.y_root)

    def _delete_selected_ref(self, event=None):
        """Supprime la référence sélectionnée."""
        selection = self.refs_tree.selection()
        if not selection:
            return

        item_id = selection[0]
        ref_id = self.refs_tree.item(item_id, "tags")
        if not ref_id:
            return
        ref_id = ref_id[0]

        # Confirmer la suppression
        if not messagebox.askyesno(
            "Confirmer la suppression",
            "Voulez-vous vraiment supprimer cette référence de codage ?"
        ):
            return

        # Supprimer la référence
        ref = CodeReference.get(self.project.db, ref_id)
        if ref:
            ref.delete(self.project.db)

            # Rafraîchir les affichages
            self.display_node_references()
            if self.current_source:
                self.refresh_document_codes()
                self.display_source()
            self.refresh_nodes()

            self.update_status("Référence supprimée")

    # ===== Handlers pour les codages du document (doc_codes_tree) =====

    def _on_doc_code_double_click(self, event):
        """Double-clic sur un codage : naviguer vers la position dans le texte."""
        selection = self.doc_codes_tree.selection()
        if not selection:
            return

        item_id = selection[0]
        # L'ID de l'item est l'ID de la référence
        ref_id = item_id

        # Récupérer la référence
        ref = CodeReference.get(self.project.db, ref_id)
        if not ref:
            return

        # Naviguer vers la position
        self._scroll_to_position(ref.start_pos, ref.end_pos)

    def _show_doc_codes_context_menu(self, event):
        """Affiche le menu contextuel pour les codages du document."""
        # Sélectionner l'item sous le curseur
        item = self.doc_codes_tree.identify_row(event.y)
        if item:
            self.doc_codes_tree.selection_set(item)

        self.doc_codes_context_menu.delete(0, tk.END)

        if item:
            self.doc_codes_context_menu.add_command(
                label="📍 Aller à cette position",
                command=lambda: self._on_doc_code_double_click(None)
            )
            self.doc_codes_context_menu.add_separator()
            self.doc_codes_context_menu.add_command(
                label="🗑️ Supprimer ce codage",
                command=self._delete_selected_doc_code
            )

        self.doc_codes_context_menu.tk_popup(event.x_root, event.y_root)

    def _delete_selected_doc_code(self, event=None):
        """Supprime le codage sélectionné dans le document."""
        selection = self.doc_codes_tree.selection()
        if not selection:
            return

        item_id = selection[0]
        ref_id = item_id

        # Confirmer la suppression
        if not messagebox.askyesno(
            "Confirmer la suppression",
            "Voulez-vous vraiment supprimer ce codage ?"
        ):
            return

        # Supprimer la référence
        ref = CodeReference.get(self.project.db, ref_id)
        if ref:
            ref.delete(self.project.db)

            # Rafraîchir les affichages
            self.refresh_document_codes()
            if self.current_source:
                self.display_source()
            self.refresh_nodes()
            if self.selected_node:
                self.display_node_references()

            self.update_status("Codage supprimé")

    def _get_reference_at_position(self, index: str) -> CodeReference | None:
        """Trouve une référence de codage à la position donnée dans le texte.

        Args:
            index: Index Tkinter (ex: "1.0", "2.5")

        Returns:
            CodeReference si trouvée, None sinon
        """
        if not self.project or not self.current_source:
            return None

        # Convertir l'index Tkinter en position caractère
        char_count = self.content_text.count("1.0", index, "chars")
        if not char_count:
            return None
        char_pos = char_count[0]

        # Chercher parmi les références de ce document
        refs = CodeReference.get_by_source(self.project.db, self.current_source.id)
        for ref in refs:
            if ref.start_pos <= char_pos < ref.end_pos:
                return ref

        return None

    def _change_reference_node(self, ref: CodeReference):
        """Affiche un dialogue pour changer le nœud d'une référence."""
        if not self.project:
            return

        dialog = tk.Toplevel(self.root)
        dialog.title("Changer le noeud")
        dialog.geometry("350x400")
        dialog.transient(self.root)
        dialog.grab_set()

        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Texte de la référence
        content_preview = ref.content[:50] + "..." if len(ref.content) > 50 else ref.content
        ttk.Label(
            main_frame,
            text=f"Texte: \"{content_preview}\"",
            wraplength=300
        ).pack(anchor=tk.W)

        ttk.Label(main_frame, text="Choisissez le nouveau noeud:").pack(
            pady=(10, 5), anchor=tk.W
        )

        # Liste des nœuds
        list_frame = ttk.Frame(main_frame)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        node_list = tk.Listbox(list_frame, yscrollcommand=scrollbar.set)
        node_list.pack(fill=tk.BOTH, expand=True)
        scrollbar.config(command=node_list.yview)

        nodes = Node.get_all(self.project.db)
        for node in nodes:
            node_list.insert(tk.END, f"● {node.name}")
            # Marquer le nœud actuel
            if node.id == ref.node_id:
                node_list.itemconfig(tk.END, foreground="gray")

        def apply_change():
            selection = node_list.curselection()
            if not selection:
                messagebox.showwarning("Attention", "Veuillez sélectionner un nœud.")
                return

            selected_node = nodes[selection[0]]
            if selected_node.id == ref.node_id:
                dialog.destroy()
                return

            # Mettre à jour la référence
            ref.node_id = selected_node.id
            ref.save(self.project.db)

            # Rafraîchir les affichages
            self.refresh_document_codes()
            if self.current_source:
                self.display_source()
            self.refresh_nodes()
            if self.selected_node:
                self.display_node_references()

            self.update_status(f"Reference deplacee vers '{selected_node.name}'")
            dialog.destroy()

        ttk.Button(main_frame, text="Appliquer", command=apply_change).pack(pady=(15, 0))

    def _delete_reference(self, ref: CodeReference):
        """Supprime une référence de codage après confirmation."""
        if not self.project:
            return

        # Confirmer la suppression
        node = Node.get(self.project.db, ref.node_id)
        node_name = node.name if node else "inconnu"
        content_preview = ref.content[:30] + "..." if len(ref.content) > 30 else ref.content

        if not messagebox.askyesno(
            "Confirmer la suppression",
            f"Supprimer la référence '{node_name}' ?\n\nTexte: \"{content_preview}\""
        ):
            return

        # Supprimer
        ref.delete(self.project.db)

        # Rafraîchir les affichages
        self.refresh_document_codes()
        if self.current_source:
            self.display_source()
        self.refresh_nodes()
        if self.selected_node:
            self.display_node_references()

        self.update_status("Référence supprimée")

    def _show_all_nodes_for_coding(self):
        """Affiche un dialogue pour selectionner un noeud parmi tous."""
        if not self.project:
            return

        dialog = tk.Toplevel(self.root)
        dialog.title("Selectionner un noeud")
        dialog.geometry("350x400")
        dialog.transient(self.root)
        dialog.grab_set()

        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(main_frame, text="Choisissez un noeud pour coder la selection:").pack(
            anchor=tk.W
        )

        # Liste des noeuds
        tree_frame = ttk.Frame(main_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0))

        tree = ttk.Treeview(tree_frame, show="tree", selectmode="browse")
        scroll = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscrollcommand=scroll.set)

        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)

        nodes = Node.get_all(self.project.db)
        for node in nodes:
            tree.tag_configure(node.color, foreground=node.color)
            tree.insert("", tk.END, iid=node.id, text=f"* {node.name}", tags=(node.color,))

        def on_select():
            selection = tree.selection()
            if selection:
                node_id = selection[0]
                node = Node.get_by_id(self.project.db, node_id)
                if node:
                    self._code_with_node(node)
            dialog.destroy()

        ttk.Button(main_frame, text="Coder", command=on_select).pack(pady=(15, 0))

        tree.bind("<Double-1>", lambda e: on_select())

    def _search_selected_text(self, text: str):
        """Recherche le texte sélectionné dans les sources."""
        # Implémenter la recherche
        self.update_status(f"Recherche: {text[:50]}...")

    def _rename_node(self):
        """Renomme le noeud selectionne."""
        if not self.selected_node:
            return

        dialog = tk.Toplevel(self.root)
        dialog.title("Renommer le noeud")
        dialog.geometry("350x130")
        dialog.transient(self.root)
        dialog.grab_set()

        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(main_frame, text="Nouveau nom:").pack(anchor=tk.W)
        name_var = tk.StringVar(value=self.selected_node.name)
        entry = ttk.Entry(main_frame, textvariable=name_var, width=40)
        entry.pack(pady=5, fill=tk.X)
        entry.focus_set()
        entry.select_range(0, tk.END)

        def save():
            new_name = name_var.get().strip()
            if new_name and new_name != self.selected_node.name:
                self.selected_node.name = new_name
                self.selected_node.save(self.project.db)
                self.refresh_nodes()
                self.update_status(f"Noeud renomme en '{new_name}'")
            dialog.destroy()

        dialog.bind("<Return>", lambda e: save())
        dialog.bind("<Escape>", lambda e: dialog.destroy())

        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=(15, 0))
        ttk.Button(btn_frame, text="Annuler", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Enregistrer", command=save).pack(side=tk.LEFT)

    def _change_node_color(self):
        """Change la couleur du noeud selectionne."""
        if not self.selected_node:
            return

        dialog = tk.Toplevel(self.root)
        dialog.title("Couleur du noeud")
        dialog.geometry("350x160")
        dialog.transient(self.root)
        dialog.grab_set()

        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(
            main_frame,
            text=f"Couleur pour '{self.selected_node.name}':",
        ).pack(anchor=tk.W)

        color_var = tk.StringVar(value=self.selected_node.color)

        color_frame = ttk.Frame(main_frame)
        color_frame.pack(pady=10)

        colors = [
            "#e74c3c", "#e91e63", "#9b59b6", "#673ab7",
            "#3498db", "#2196f3", "#00bcd4", "#009688",
            "#2ecc71", "#4caf50", "#8bc34a", "#cddc39",
            "#f39c12", "#ff9800", "#ff5722", "#795548",
        ]

        for i, color in enumerate(colors):
            btn = tk.Button(
                color_frame,
                bg=color,
                activebackground=color,
                width=2,
                height=1,
                relief=tk.FLAT if color != self.selected_node.color else tk.SUNKEN,
                cursor="hand2",
                command=lambda c=color: color_var.set(c),
            )
            btn.grid(row=i // 8, column=i % 8, padx=2, pady=2)

        def save():
            new_color = color_var.get()
            if new_color != self.selected_node.color:
                self.selected_node.color = new_color
                self.selected_node.save(self.project.db)
                self.refresh_nodes()
                if self.current_source:
                    self.display_source()
                self.update_status("Couleur mise a jour")
            dialog.destroy()

        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(pady=(10, 0))
        ttk.Button(btn_frame, text="Annuler", command=dialog.destroy).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Appliquer", command=save).pack(side=tk.LEFT)

    def _quick_code_from_selection(self):
        """Crée un nœud depuis la sélection et code immédiatement (Ctrl+Shift+K)."""
        if not self.project or not self.current_source:
            return

        try:
            selected_text = self.content_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            if selected_text.strip():
                short_text = selected_text[:30].replace("\n", " ").strip()
                if len(selected_text) > 30:
                    short_text += "..."
                self.create_node(initial_name=short_text, code_selection=True)
        except tk.TclError:
            messagebox.showinfo("Information", "Sélectionnez du texte à coder.")

    def _delete_if_node_focused(self):
        """Supprime le nœud si le focus est sur l'arbre des nœuds."""
        if self.root.focus_get() == self.nodes_tree and self.selected_node:
            self.delete_node()

    def auto_detect_nodes(self):
        """Lance la détection automatique de nœuds."""
        if not self.project:
            messagebox.showwarning(
                "Attention",
                "Veuillez d'abord créer ou ouvrir un projet."
            )
            return

        # Récupérer les sources avec contenu textuel
        all_sources = Source.get_all(self.project.db)
        text_sources = [
            {
                "id": s.id,
                "name": s.name,
                "type": s.type.value,
                "content": s.content,
            }
            for s in all_sources
            if s.content and len(s.content) > 100
        ]

        if not text_sources:
            messagebox.showwarning(
                "Attention",
                "Aucune source avec contenu textuel suffisant.\n"
                "Importez des documents avant d'utiliser la détection automatique."
            )
            return

        # Récupérer les nœuds existants
        existing_nodes = [
            {"id": n.id, "name": n.name}
            for n in Node.get_all(self.project.db)
        ]

        # Charger les settings
        settings = self.settings_manager.settings
        saved_settings = {
            "llm_provider": settings.llm_provider,
            "llm_model": settings.llm_model,
            "max_themes": settings.autocoding_max_themes,
            "min_cluster_size": settings.autocoding_min_cluster_size,
            "confidence_threshold": settings.autocoding_confidence_threshold,
        }

        # Ouvrir le dialogue de configuration
        config_dialog = AutoCodingConfigDialog(
            self.root,
            sources=text_sources,
            existing_nodes=existing_nodes,
            settings=saved_settings,
        )
        self.root.wait_window(config_dialog)

        if config_dialog.cancelled or not config_dialog.result_config:
            return

        config = config_dialog.result_config
        sources_to_analyze = config_dialog.result_sources

        # Sauvegarder les paramètres
        settings.llm_provider = config.llm_provider.value
        settings.llm_model = config.llm_model
        settings.autocoding_max_themes = config.max_themes
        settings.autocoding_min_cluster_size = config.min_cluster_size
        settings.autocoding_confidence_threshold = config.confidence_threshold
        self.settings_manager.save()

        # Lancer l'analyse dans un thread
        self._run_auto_coding(sources_to_analyze, config, existing_nodes)

    def _run_auto_coding(self, sources: list, config, existing_nodes: list):
        """Exécute l'auto-codage de manière asynchrone."""
        from ..analysis.auto_coding import AutoCodingEngine, create_nodes_from_proposals

        # Dialogue de progression
        progress_dialog = AutoCodingProgressDialog(self.root, len(sources))

        def do_analysis():
            try:
                # Créer le moteur
                settings = self.settings_manager.settings
                engine = AutoCodingEngine(
                    embedding_model=settings.autocoding_embedding_model,
                    llm_provider=config.llm_provider,
                    llm_model=config.llm_model,
                    ollama_url=settings.ollama_url,
                )

                # Callback de progression
                def on_progress(progress: float, message: str):
                    if not progress_dialog.cancelled:
                        progress_dialog.after(
                            0,
                            lambda: progress_dialog.update_progress(progress, message)
                        )

                # Exécuter l'analyse
                result = engine.analyze(
                    sources=sources,
                    config=config,
                    existing_nodes=existing_nodes,
                    progress_callback=on_progress,
                )

                # Fermer le dialogue de progression et afficher les résultats
                self.root.after(0, lambda: self._show_auto_coding_results(result, progress_dialog))

            except Exception as e:
                logger.error(f"Erreur auto-codage: {e}")
                self.root.after(0, lambda: self._on_auto_coding_error(str(e), progress_dialog))

        thread = threading.Thread(target=do_analysis, daemon=True)
        thread.start()

    def _show_auto_coding_results(self, result, progress_dialog):
        """Affiche les résultats de l'auto-codage."""
        from ..analysis.auto_coding import create_nodes_from_proposals

        progress_dialog.complete()

        if not result.proposals:
            messagebox.showinfo(
                "Résultat",
                "Aucun thème détecté.\n\n"
                "Essayez avec des paramètres différents ou ajoutez plus de contenu."
            )
            return

        # Afficher le dialogue de preview
        preview_dialog = AutoCodingPreviewDialog(self.root, result)
        self.root.wait_window(preview_dialog)

        if not preview_dialog.approved:
            return

        # Créer les nœuds et les codages
        try:
            created = create_nodes_from_proposals(
                self.project.db,
                result.selected_proposals,
                parent_id=None,
            )

            self.refresh_nodes()
            self.refresh_sources()

            n_nodes = len([c for c in created if c.get("is_new")])
            n_segments = result.total_selected_segments

            messagebox.showinfo(
                "Succès",
                f"Auto-codage terminé !\n\n"
                f"• {n_nodes} nouveau(x) nœud(s) créé(s)\n"
                f"• {n_segments} segment(s) codé(s)"
            )

            self.update_status(f"Auto-codage: {n_nodes} nœuds, {n_segments} codages")

        except Exception as e:
            logger.error(f"Erreur création nœuds: {e}")
            messagebox.showerror(
                "Erreur",
                f"Erreur lors de la création des nœuds:\n{e}"
            )

    def _on_auto_coding_error(self, error_message: str, progress_dialog):
        """Gère une erreur d'auto-codage."""
        progress_dialog.complete()
        messagebox.showerror(
            "Erreur",
            f"Erreur lors de l'analyse automatique:\n\n{error_message}\n\n"
            "Vérifiez que les dépendances sont installées:\n"
            "pip install sentence-transformers umap-learn hdbscan"
        )

    def delete_node(self):
        """Supprime le nœud sélectionné."""
        if not self.project or not self.selected_node:
            return

        if messagebox.askyesno(
            "Confirmer",
            f"Supprimer le nœud '{self.selected_node.name}' et ses références?",
        ):
            self.selected_node.delete(self.project.db)
            self.selected_node = None
            self.refresh_nodes()

    def code_selection(self):
        """Code la sélection avec le nœud sélectionné."""
        if not self.project or not self.current_source or not self.selected_node:
            messagebox.showinfo(
                "Information",
                "Sélectionnez une source et un nœud, puis sélectionnez du texte à coder.",
            )
            return

        try:
            sel_start = self.content_text.index(tk.SEL_FIRST)
            sel_end = self.content_text.index(tk.SEL_LAST)
            selected_text = self.content_text.get(sel_start, sel_end)

            # Convertir les indices tkinter en positions
            start_count = self.content_text.count("1.0", sel_start, "chars")
            end_count = self.content_text.count("1.0", sel_end, "chars")
            start_pos = int(start_count[0]) if start_count else 0
            end_pos = int(end_count[0]) if end_count else 0

            # Créer la référence de codage
            ref = CodeReference(
                node_id=self.selected_node.id,
                source_id=self.current_source.id,
                start_pos=start_pos,
                end_pos=end_pos,
                content=selected_text,
            )
            ref.save(self.project.db)

            # Surligner le texte codé
            self.highlight_coding(sel_start, sel_end, self.selected_node.color)
            self.refresh_document_codes()
            self.update_status(f"Texte codé avec '{self.selected_node.name}'")

        except tk.TclError:
            messagebox.showinfo("Information", "Veuillez sélectionner du texte à coder.")

    def highlight_coding(self, start, end, color):
        """Surligne un passage codé."""
        tag_name = f"code_{color}"
        # Tkinter ne supporte pas les couleurs avec alpha, créer une version claire
        light_color = self._lighten_color(color, factor=0.7)
        self.content_text.tag_configure(tag_name, background=light_color)
        self.content_text.tag_add(tag_name, start, end)

    def _lighten_color(self, hex_color: str, factor: float = 0.7) -> str:
        """Éclaircit une couleur hex pour le surlignage.

        Args:
            hex_color: Couleur au format #RRGGBB
            factor: Facteur d'éclaircissement (0-1, plus haut = plus clair)

        Returns:
            Couleur éclaircie au format #RRGGBB
        """
        # Nettoyer la couleur
        color = hex_color.lstrip('#')
        if len(color) != 6:
            return "#FFFF99"  # Jaune par défaut si format invalide

        try:
            r = int(color[0:2], 16)
            g = int(color[2:4], 16)
            b = int(color[4:6], 16)

            # Mélanger avec du blanc
            r = int(r + (255 - r) * factor)
            g = int(g + (255 - g) * factor)
            b = int(b + (255 - b) * factor)

            return f"#{r:02x}{g:02x}{b:02x}"
        except ValueError:
            return "#FFFF99"

    # --- Actions d'analyse ---

    def show_search(self):
        """Affiche la recherche."""
        self.analysis_notebook.select(self.search_frame)

        dialog = tk.Toplevel(self.root)
        dialog.title("Recherche")
        dialog.geometry("400x150")
        dialog.transient(self.root)

        main_frame = ttk.Frame(dialog, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(main_frame, text="Rechercher:").pack(anchor=tk.W)
        search_var = tk.StringVar()
        ttk.Entry(main_frame, textvariable=search_var, width=40).pack(pady=5, fill=tk.X)

        def search():
            query = search_var.get().strip()
            if query and self.project:
                from ..analysis.search import SearchEngine

                engine = SearchEngine(self.project.db)
                results = engine.search(query)

                self.search_results.delete(*self.search_results.get_children())
                for result in results:
                    self.search_results.insert(
                        "",
                        tk.END,
                        values=(result.item_type, result.snippet[:100]),
                    )

                dialog.destroy()
                self.update_status(f"{len(results)} resultat(s) trouve(s)")

        ttk.Button(main_frame, text="Rechercher", command=search).pack(pady=(15, 0))

    def quick_search(self):
        """Recherche rapide."""
        query = self.quick_search_var.get().strip()
        if query and self.project:
            from ..analysis.search import SearchEngine

            engine = SearchEngine(self.project.db)
            results = engine.search(query, limit=20)

            self.search_results.delete(*self.search_results.get_children())
            for result in results:
                self.search_results.insert(
                    "",
                    tk.END,
                    values=(result.item_type, result.snippet[:100]),
                )

            self.analysis_notebook.select(self.search_frame)
            self.update_status(f"{len(results)} résultat(s)")

    def show_wordcloud(self):
        """Génère et affiche un nuage de mots."""
        if not self.project:
            return

        from ..visualization.wordcloud_viz import WordCloudGenerator

        generator = WordCloudGenerator(self.project.db)
        image_data = generator.generate()

        if image_data:
            # Afficher dans une nouvelle fenêtre
            window = tk.Toplevel(self.root)
            window.title("Nuage de mots")
            window.geometry("850x500")

            from PIL import Image, ImageTk
            import io

            img = Image.open(io.BytesIO(image_data))
            photo = ImageTk.PhotoImage(img)

            label = ttk.Label(window, image=photo)
            label.image = photo
            label.pack(expand=True)

            ttk.Button(
                window,
                text="Sauvegarder",
                command=lambda: self.save_visualization(image_data, "wordcloud.png"),
            ).pack(pady=10)

    def show_mindmap(self):
        """Génère et affiche une carte mentale."""
        if not self.project:
            return

        from ..visualization.mindmap import MindMapGenerator

        generator = MindMapGenerator(self.project.db)
        image_data = generator.generate()

        if image_data:
            window = tk.Toplevel(self.root)
            window.title("Carte mentale")
            window.geometry("1250x850")

            from PIL import Image, ImageTk
            import io

            img = Image.open(io.BytesIO(image_data))
            photo = ImageTk.PhotoImage(img)

            label = ttk.Label(window, image=photo)
            label.image = photo
            label.pack(expand=True)

            ttk.Button(
                window,
                text="Sauvegarder",
                command=lambda: self.save_visualization(image_data, "mindmap.png"),
            ).pack(pady=10)

    def show_sociogram(self):
        """Génère et affiche un sociogramme."""
        if not self.project:
            return

        from ..visualization.sociogram import SociogramGenerator

        generator = SociogramGenerator(self.project.db)
        image_data = generator.generate_node_cooccurrence()

        if image_data:
            window = tk.Toplevel(self.root)
            window.title("Sociogramme")
            window.geometry("1050x850")

            from PIL import Image, ImageTk
            import io

            img = Image.open(io.BytesIO(image_data))
            photo = ImageTk.PhotoImage(img)

            label = ttk.Label(window, image=photo)
            label.image = photo
            label.pack(expand=True)

            ttk.Button(
                window,
                text="Sauvegarder",
                command=lambda: self.save_visualization(image_data, "sociogram.png"),
            ).pack(pady=10)

    def save_visualization(self, image_data: bytes, default_name: str):
        """Sauvegarde une visualisation."""
        path = filedialog.asksaveasfilename(
            defaultextension=".png",
            initialfile=default_name,
            filetypes=[("PNG", "*.png"), ("Tous", "*.*")],
        )
        if path:
            Path(path).write_bytes(image_data)
            self.update_status(f"Image sauvegardée: {path}")

    # --- Actions diverses ---

    def delete_source(self):
        """Supprime la source sélectionnée."""
        if not self.project or not self.current_source:
            return

        if messagebox.askyesno(
            "Confirmer",
            f"Supprimer la source '{self.current_source.name}'?",
        ):
            self.current_source.delete(self.project.db)
            self.current_source = None
            self.refresh_sources()
            self.clear_content()

    # ===== Menu contextuel des sources =====

    def _show_source_context_menu(self, event):
        """Affiche le menu contextuel pour les sources."""
        # Sélectionner l'item sous le curseur
        item = self.sources_tree.identify_row(event.y)
        if item:
            self.sources_tree.selection_set(item)
            # Mettre à jour current_source
            source_id = item
            self.current_source = Source.get(self.project.db, source_id)

        self.source_context_menu.delete(0, tk.END)

        if self.current_source:
            # Ouvrir
            self.source_context_menu.add_command(
                label="📄 Ouvrir",
                command=lambda: self.display_source()
            )

            # Export transcription pour audio/vidéo
            if self.current_source.type in (SourceType.AUDIO, SourceType.VIDEO):
                self.source_context_menu.add_separator()
                self.source_context_menu.add_command(
                    label="📝 Exporter la transcription (.txt)",
                    command=lambda: self._export_transcription("txt")
                )
                self.source_context_menu.add_command(
                    label="📄 Exporter la transcription (.docx)",
                    command=lambda: self._export_transcription("docx")
                )

            self.source_context_menu.add_separator()
            self.source_context_menu.add_command(
                label="🗑️ Supprimer",
                command=self.delete_source
            )

        self.source_context_menu.tk_popup(event.x_root, event.y_root)

    def _export_transcription(self, format: str):
        """Exporte la transcription d'une source audio/vidéo."""
        if not self.current_source:
            return

        if not self.current_source.content:
            messagebox.showwarning(
                "Aucune transcription",
                "Cette source n'a pas de transcription à exporter."
            )
            return

        # Nom par défaut
        default_name = Path(self.current_source.name).stem + "_transcription"

        if format == "txt":
            file_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                initialfile=default_name + ".txt",
                filetypes=[("Fichiers texte", "*.txt"), ("Tous les fichiers", "*.*")]
            )
            if file_path:
                Path(file_path).write_text(self.current_source.content, encoding="utf-8")
                self.update_status(f"Transcription exportée: {file_path}")

        elif format == "docx":
            try:
                from docx import Document
                from docx.shared import Pt
            except ImportError:
                messagebox.showerror(
                    "Module manquant",
                    "python-docx n'est pas installé.\n"
                    "Installez-le avec: pip install python-docx"
                )
                return

            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                initialfile=default_name + ".docx",
                filetypes=[("Document Word", "*.docx"), ("Tous les fichiers", "*.*")]
            )
            if file_path:
                doc = Document()
                doc.add_heading(f"Transcription: {self.current_source.name}", level=1)

                # Ajouter les métadonnées si disponibles
                metadata = self.current_source.metadata or {}
                if "duration" in metadata:
                    duration = metadata["duration"]
                    minutes = int(duration // 60)
                    seconds = int(duration % 60)
                    doc.add_paragraph(f"Durée: {minutes}:{seconds:02d}")

                transcription_meta = metadata.get("transcription", {})
                if transcription_meta.get("model"):
                    doc.add_paragraph(f"Modèle Whisper: {transcription_meta['model']}")
                if transcription_meta.get("language_detected"):
                    doc.add_paragraph(f"Langue détectée: {transcription_meta['language_detected']}")

                doc.add_paragraph("")  # Ligne vide

                # Contenu
                for paragraph_text in self.current_source.content.split("\n\n"):
                    if paragraph_text.strip():
                        p = doc.add_paragraph(paragraph_text.strip())
                        p.style.font.size = Pt(11)

                doc.save(file_path)
                self.update_status(f"Transcription exportée: {file_path}")

    def show_about(self):
        """Affiche la fenêtre À propos."""
        messagebox.showinfo(
            "À propos",
            f"Lele - Analyse Qualitative de Données\n\n"
            f"Version {__version__}\n\n"
            "Application d'analyse qualitative inspirée de NVivo.\n"
            "Supporte l'import de multiples formats, le codage,\n"
            "et diverses visualisations.",
        )

    def show_help(self):
        """Affiche le guide d'utilisation."""
        dialog = tk.Toplevel(self.root)
        dialog.title("Guide d'utilisation - Lele")
        dialog.geometry("450x650")
        dialog.transient(self.root)

        # Frame principale avec scrollbar et padding
        main_frame = ttk.Frame(dialog, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Obtenir la couleur de fond du style ttk (coherent avec les autres dialogues)
        style = ttk.Style()
        bg_color = style.lookup("TFrame", "background") or "#f0f0f0"

        # Canvas et scrollbar
        canvas = tk.Canvas(main_frame, highlightthickness=0, bg=bg_color)
        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=canvas.yview)
        scrollable_frame = tk.Frame(canvas, bg=bg_color)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Contenu de l'aide
        help_sections = [
            ("Demarrage rapide", [
                "1. Creez un projet: Fichier > Nouveau projet (Ctrl+N)",
                "2. Importez des fichiers: glissez-deposez ou Fichier > Importer",
                "3. Creez des noeuds (codes) dans le panneau Noeuds",
                "4. Selectionnez du texte et codez-le avec Ctrl+K",
            ]),
            ("Gestion des projets", [
                "- Nouveau projet: Fichier > Nouveau projet",
                "- Ouvrir: Fichier > Ouvrir projet",
                "- Sauvegarder: Fichier > Sauvegarder (Ctrl+S)",
                "- Projets recents accessibles via Fichier > Projets recents",
            ]),
            ("Import de donnees", [
                "Formats supportes:",
                "- Documents: TXT, PDF, Word (.docx), RTF",
                "- Media: Audio (MP3, WAV, FLAC), Video (MP4, AVI, MOV)",
                "- Images: JPG, PNG, GIF, BMP, TIFF",
                "- Tableurs: Excel (.xlsx), CSV",
                "- Bibliographie: RIS, BibTeX",
                "",
                "Pour l'audio/video, la transcription automatique est disponible",
                "via le moteur Whisper (OpenAI).",
            ]),
            ("Codage", [
                "Methode classique:",
                "1. Ouvrez une source (double-clic)",
                "2. Selectionnez un noeud dans le panneau Noeuds",
                "3. Surlignez le texte a coder",
                "4. Appuyez sur Ctrl+K",
                "",
                "Methode rapide:",
                "- Surlignez du texte > clic droit > \"Coder avec\"",
                "- Ou double-cliquez sur un noeud avec du texte selectionne",
                "",
                "Creer et coder en meme temps:",
                "- Surlignez du texte > Ctrl+Shift+K",
            ]),
            ("Detection automatique de themes", [
                "1. Cliquez sur \"Auto\" dans le panneau Noeuds",
                "2. Selectionnez les sources a analyser",
                "3. Configurez les parametres (granularite, nb max de themes)",
                "4. Validez les themes proposes avant creation",
                "",
                "Pour un meilleur nommage, configurez Ollama dans",
                "Parametres > IA / LLM local.",
            ]),
            ("Visualisations", [
                "Menu Analyse propose:",
                "- Nuage de mots: frequence des termes",
                "- Carte mentale: hierarchie des noeuds",
                "- Sociogramme: co-occurrences entre noeuds",
                "- Matrice: croisement noeuds x sources",
            ]),
            ("Conseils", [
                "- Utilisez les couleurs pour distinguer les themes",
                "- Creez une hierarchie de noeuds (sous-noeuds)",
                "- Exportez vos transcriptions en .docx pour les partager",
                "- Sauvegardez regulierement votre projet",
            ]),
        ]

        for title, lines in help_sections:
            # Titre de section
            title_label = tk.Label(
                scrollable_frame,
                text=title,
                font=("", 12, "bold"),
                bg=bg_color,
            )
            title_label.pack(anchor=tk.W, pady=(15, 5), padx=10)

            # Contenu
            for line in lines:
                line_label = tk.Label(
                    scrollable_frame,
                    text=line,
                    wraplength=500,
                    justify=tk.LEFT,
                    bg=bg_color,
                )
                line_label.pack(anchor=tk.W, padx=20)

        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Binding molette souris avec nettoyage à la fermeture
        mousewheel_bound = [False]  # Liste pour pouvoir modifier dans les closures

        def _on_mousewheel(event):
            try:
                if canvas.winfo_exists():
                    canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
            except tk.TclError:
                pass

        def _bind_mousewheel(event):
            if not mousewheel_bound[0]:
                dialog.bind_all("<MouseWheel>", _on_mousewheel)
                mousewheel_bound[0] = True

        def _unbind_mousewheel(event=None):
            if mousewheel_bound[0]:
                try:
                    dialog.unbind_all("<MouseWheel>")
                except tk.TclError:
                    pass
                mousewheel_bound[0] = False

        canvas.bind("<Enter>", _bind_mousewheel)
        canvas.bind("<Leave>", _unbind_mousewheel)
        scrollable_frame.bind("<Enter>", _bind_mousewheel)
        scrollable_frame.bind("<Leave>", _unbind_mousewheel)

        def _on_close():
            _unbind_mousewheel()
            dialog.destroy()

        # Bouton fermer avec nettoyage
        ttk.Button(dialog, text="Fermer", command=_on_close).pack(pady=10)
        dialog.protocol("WM_DELETE_WINDOW", _on_close)

    def show_shortcuts_help(self):
        """Affiche la liste des raccourcis clavier."""
        dialog = tk.Toplevel(self.root)
        dialog.title("Raccourcis clavier - Lele")
        dialog.geometry("500x500")
        dialog.transient(self.root)

        # Frame principale avec padding
        outer_frame = ttk.Frame(dialog, padding="20")
        outer_frame.pack(fill=tk.BOTH, expand=True)

        # Titre
        ttk.Label(
            outer_frame,
            text="⌨️ Raccourcis clavier",
            font=("", 14, "bold"),
        ).pack(pady=(0, 10))

        # Frame avec scrollbar
        main_frame = ttk.Frame(outer_frame)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Treeview pour les raccourcis
        tree = ttk.Treeview(
            main_frame,
            columns=("action",),
            show="headings",
            height=18,
        )
        tree.heading("#0", text="")
        tree.heading("action", text="Action")
        tree.column("action", width=350)

        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscrollcommand=scrollbar.set)

        shortcuts = [
            ("", "── Fichier ──"),
            ("Ctrl+N", "Nouveau projet"),
            ("Ctrl+O", "Ouvrir projet"),
            ("Ctrl+S", "Sauvegarder"),
            ("Ctrl+Q", "Quitter"),
            ("", ""),
            ("", "── Édition ──"),
            ("Ctrl+Z", "Annuler"),
            ("Ctrl+Y", "Rétablir"),
            ("Ctrl+F", "Rechercher"),
            ("", ""),
            ("", "── Codage ──"),
            ("Ctrl+K", "Coder la sélection avec le nœud actif"),
            ("Ctrl+Shift+N", "Créer un nouveau nœud"),
            ("Ctrl+Shift+K", "Créer un nœud depuis la sélection et coder"),
            ("F2", "Renommer le nœud sélectionné"),
            ("Suppr", "Supprimer le nœud sélectionné"),
            ("", ""),
            ("", "── Aide ──"),
            ("F1", "Afficher l'aide"),
            ("", ""),
            ("", "── Actions rapides ──"),
            ("Double-clic nœud", "Coder la sélection avec ce nœud"),
            ("Clic droit texte", "Menu contextuel de codage"),
            ("Clic droit nœud", "Menu contextuel du nœud"),
        ]

        for shortcut, action in shortcuts:
            if shortcut == "" and action.startswith("──"):
                # Ligne de section
                tree.insert("", tk.END, values=(action,), tags=("section",))
            elif shortcut == "":
                # Ligne vide
                tree.insert("", tk.END, values=("",))
            else:
                tree.insert("", tk.END, values=(f"{shortcut}  →  {action}",))

        tree.tag_configure("section", font=("", 10, "bold"))

        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Bouton fermer
        ttk.Button(outer_frame, text="Fermer", command=dialog.destroy).pack(pady=(15, 0))

    def show_transcription_settings(self):
        """Affiche les paramètres de transcription."""
        dialog = TranscriptionSettingsDialog(
            self.root,
            current_model=self.whisper_model,
            current_language=self.whisper_language,
            show_transcribe_option=False,
        )
        self.root.wait_window(dialog)

        if not dialog.cancelled:
            self.whisper_model = dialog.result_model
            self.whisper_language = dialog.result_language

            # Sauvegarder dans les settings
            self.settings_manager.settings.whisper_model = self.whisper_model
            self.settings_manager.settings.whisper_language = self.whisper_language
            self.settings_manager.save()

            lang_text = self.whisper_language or "auto"
            self.update_status(
                f"Paramètres de transcription: modèle={self.whisper_model}, langue={lang_text}"
            )

    def show_llm_settings(self):
        """Affiche les paramètres IA / LLM."""
        settings = self.settings_manager.settings

        dialog = LLMSettingsDialog(
            self.root,
            llm_provider=settings.llm_provider,
            llm_model=settings.llm_model,
            ollama_url=settings.ollama_url,
            embedding_model=settings.autocoding_embedding_model,
        )
        self.root.wait_window(dialog)

        if not dialog.cancelled:
            # Sauvegarder les paramètres
            settings.llm_provider = dialog.result_provider
            settings.llm_model = dialog.result_model
            settings.ollama_url = dialog.result_ollama_url
            settings.autocoding_embedding_model = dialog.result_embedding_model
            self.settings_manager.save()

            self.update_status(
                f"Paramètres IA: {dialog.result_provider}/{dialog.result_model}"
            )

    # --- Rafraîchissement de l'interface ---

    def refresh_all(self):
        """Rafraîchit toute l'interface."""
        self.refresh_sources()
        self.refresh_nodes()

    def refresh_sources(self):
        """Rafraîchit la liste des sources."""
        self.sources_tree.delete(*self.sources_tree.get_children())

        if not self.project:
            return

        filter_type = self.source_type_filter.get()
        type_map = {
            "Texte": SourceType.TEXT,
            "Audio": SourceType.AUDIO,
            "Vidéo": SourceType.VIDEO,
            "Image": SourceType.IMAGE,
            "Tableur": SourceType.SPREADSHEET,
            "PDF": SourceType.PDF,
        }

        source_filter = type_map.get(filter_type)
        sources = Source.get_all(self.project.db, source_type=source_filter)

        for source in sources:
            ref_count = CodeReference.count_by_source(self.project.db, source.id)
            self.sources_tree.insert(
                "",
                tk.END,
                iid=source.id,
                text=source.name,
                values=(source.type.value, ref_count),
            )

    def refresh_nodes(self):
        """Rafraîchit l'arbre des nœuds."""
        self.nodes_tree.delete(*self.nodes_tree.get_children())

        if not self.project:
            return

        def add_nodes(parent_id, tree_parent=""):
            nodes = Node.get_all(self.project.db, parent_id=parent_id)
            for node in nodes:
                # Configurer le tag pour afficher la couleur du nœud
                tag_name = f"color_{node.color}"
                # Créer une version sombre de la couleur pour le texte lisible
                self.nodes_tree.tag_configure(tag_name, foreground=node.color)

                # Afficher avec un indicateur de couleur
                display_text = f"● {node.name}"
                self.nodes_tree.insert(
                    tree_parent,
                    tk.END,
                    iid=node.id,
                    text=display_text,
                    values=(node.reference_count,),
                    tags=(tag_name,),
                )
                # Ajouter récursivement les enfants
                add_nodes(node.id, node.id)

        add_nodes(None)

    def refresh_document_codes(self):
        """Rafraîchit la liste des codes du document actuel."""
        self.doc_codes_tree.delete(*self.doc_codes_tree.get_children())

        if not self.project or not self.current_source:
            return

        refs = CodeReference.get_by_source(self.project.db, self.current_source.id)
        for ref in refs:
            pos = f"{ref.start_pos}-{ref.end_pos}" if ref.start_pos else ""
            self.doc_codes_tree.insert(
                "",
                tk.END,
                iid=ref.id,
                text=ref.node_name,
                values=(pos,),
            )

    # --- Gestionnaires d'événements ---

    def on_source_select(self, event):
        """Gère la sélection d'une source."""
        selection = self.sources_tree.selection()
        if selection:
            source_id = selection[0]
            self.current_source = Source.get(self.project.db, source_id)
            self.display_source()

    def on_source_double_click(self, event):
        """Gère le double-clic sur une source."""
        self.on_source_select(event)

    def on_node_select(self, event):
        """Gère la sélection d'un nœud."""
        selection = self.nodes_tree.selection()
        if selection:
            node_id = selection[0]
            self.selected_node = Node.get(self.project.db, node_id)
            self.display_node_references()

    def display_source(self):
        """Affiche le contenu d'une source."""
        if not self.current_source:
            return

        # Réinitialiser le mode édition si on change de source
        if self.edit_mode:
            self._set_edit_mode(False)

        self.content_text.configure(state=tk.NORMAL)
        self.content_text.delete("1.0", tk.END)

        if self.current_source.content:
            self.content_text.insert("1.0", self.current_source.content)

            # Surligner les codes existants
            refs = CodeReference.get_by_source(self.project.db, self.current_source.id)
            for ref in refs:
                if ref.start_pos is not None and ref.end_pos is not None:
                    node = Node.get(self.project.db, ref.node_id)
                    if node:
                        start = f"1.0+{ref.start_pos}c"
                        end = f"1.0+{ref.end_pos}c"
                        self.highlight_coding(start, end, node.color)

        # Mettre en mode lecture seule par défaut
        self.content_text.configure(
            state=tk.DISABLED,
            background="#f8f8f8",
        )

        self.refresh_document_codes()
        self.update_line_numbers()

    def display_node_references(self):
        """Affiche les références d'un nœud."""
        self.refs_tree.delete(*self.refs_tree.get_children())

        if not self.selected_node or not self.project:
            return

        refs = CodeReference.get_by_node(self.project.db, self.selected_node.id)

        # Cache pour éviter de recharger le contenu de la même source plusieurs fois
        source_contents: dict[str, str] = {}

        for ref in refs:
            content = (ref.content or "")[:100]

            # Calculer le numéro de ligne
            line_num = ""
            if ref.start_pos is not None:
                # Récupérer le contenu de la source pour compter les lignes
                if ref.source_id not in source_contents:
                    source = Source.get(self.project.db, ref.source_id)
                    if source and source.content:
                        source_contents[ref.source_id] = source.content
                    else:
                        source_contents[ref.source_id] = ""

                source_content = source_contents.get(ref.source_id, "")
                if source_content and ref.start_pos <= len(source_content):
                    # Compter les sauts de ligne avant start_pos
                    line_num = source_content[:ref.start_pos].count("\n") + 1

            self.refs_tree.insert(
                "",
                tk.END,
                values=(ref.source_name, line_num, content),
                tags=(ref.id,),
            )

        self.analysis_notebook.select(self.refs_frame)

    def clear_content(self):
        """Efface la zone de contenu."""
        self.content_text.configure(state=tk.NORMAL)
        self.content_text.delete("1.0", tk.END)
        self.content_text.configure(state=tk.DISABLED)

    def update_line_numbers(self):
        """
        Met à jour les numéros de ligne en tenant compte du word-wrap.

        Chaque ligne logique peut occuper plusieurs lignes visuelles.
        On affiche le numéro uniquement sur la première ligne visuelle,
        et des espaces sur les lignes de continuation.
        """
        self.line_numbers.configure(state=tk.NORMAL)
        self.line_numbers.delete("1.0", tk.END)

        # Forcer la mise à jour de l'affichage pour avoir les bonnes dimensions
        self.content_text.update_idletasks()

        line_count = int(self.content_text.index("end-1c").split(".")[0])
        line_numbers_text = []

        for i in range(1, line_count + 1):
            # Compter combien de lignes visuelles occupe cette ligne logique
            # en utilisant count avec "displaylines"
            start_index = f"{i}.0"
            end_index = f"{i}.end"

            try:
                # Compter les lignes d'affichage entre le début et la fin de la ligne
                display_lines = self.content_text.count(start_index, end_index, "displaylines")
                if display_lines is None:
                    display_lines = 1
                else:
                    # count retourne un tuple, on prend le premier élément
                    display_lines = display_lines[0] if isinstance(display_lines, tuple) else display_lines
                    display_lines = max(1, display_lines)
            except tk.TclError:
                display_lines = 1

            # Première ligne visuelle : afficher le numéro
            line_numbers_text.append(str(i))

            # Lignes de continuation : afficher des espaces
            for _ in range(display_lines - 1):
                line_numbers_text.append("")

        self.line_numbers.insert("1.0", "\n".join(line_numbers_text))
        self.line_numbers.configure(state=tk.DISABLED)

    def update_status(self, message: str):
        """Met à jour la barre de statut."""
        self.status_label.configure(text=message)

    def run(self):
        """Lance l'application."""
        self.root.mainloop()
